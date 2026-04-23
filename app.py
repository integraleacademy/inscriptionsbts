
import os, sqlite3, json, uuid
from datetime import datetime, timedelta
from flask import Flask, render_template, request, redirect, url_for, send_file, session, abort, jsonify, flash
from werkzeug.utils import secure_filename
from utils import send_mail, send_sms_brevo, dossier_number, new_token, sign_token, make_signed_link
from dotenv import load_dotenv
from parcoursup import bp_parcoursup
from sms_templates import sms_text
from mail_templates import mail_html
from utils import get_mail_context
from utils import BTS_LABELS

# ============================
# 🔒 STARTUP INTEGRITY CHECK – STRICT (Parcoursup inclus)
# ============================
import os, re, sys, json, glob

def _read(p):
    try:
        with open(p, encoding="utf-8") as f:
            return f.read()
    except FileNotFoundError:
        return None

def run_startup_integrity_checks():
    print("\n================= 🔎 INTÉGRITÉ PROJET – DÉMARRAGE STRICT =================")
    if os.getenv("INTEGRITY_CHECK", "1") == "0":
        print("⚠️  INTEGRITY_CHECK=0 → contrôle désactivé (non recommandé en prod).")
        return

    # --- 1) FICHIERS CRITIQUES À PRÉSENCE OBLIGATOIRE ---
    DATA_DIR = os.getenv("DATA_DIR", "/data")
    required_files = [
        # Python
        "app.py",
        "parcoursup.py",
        "mail_templates.py",
        "sms_templates.py",
        # Templates & statiques
        "templates/index.html",
        "templates/admin.html",
        "templates/parcoursup.html",
        "templates/email_base.html",
        "static/js/main_front.js",
        "static/js/main_admin.js",
        "static/css/styles.css",
        "static/logo-integrale.png",
    ]
    missing = [p for p in required_files if not os.path.exists(p)]
    if not os.path.exists(DATA_DIR):
        missing.append(DATA_DIR + "/ (répertoire)")
    if missing:
        print("❌ Fichiers/répertoires manquants :")
        for m in missing: print("   •", m)
        sys.exit(1)
    print("✅ Fichiers critiques présents.")

    # --- 2) ENV REQUISES ---
    env_required = {
        "BREVO_API_KEY": os.getenv("BREVO_API_KEY"),
        "SENDER_EMAIL": os.getenv("SENDER_EMAIL", "ecole@integraleacademy.com"),
        "BASE_URL": os.getenv("BASE_URL", "https://inscriptionsbts.onrender.com"),
    }
    if not env_required["BREVO_API_KEY"]:
        print("❌ BREVO_API_KEY manquant (requis pour mails/SMS).")
        sys.exit(1)
    print(f"✅ Env OK (SENDER_EMAIL={env_required['SENDER_EMAIL']}, BASE_URL={env_required['BASE_URL']})")

    # --- 3) ROUTES FLASK DÉFINIES (tous .py) ---
    py_sources = {}
    for p in ["app.py", "parcoursup.py", "mail_templates.py", "sms_templates.py"]:
        py_sources[p] = _read(p)

    route_regex = re.compile(r"""@(?:app|bp_[a-zA-Z0-9_]+)\.route\(\s*['"]([^'"]+)['"]""")
    defined_routes = set()
    for name, src in py_sources.items():
        if not src: continue
        for r in route_regex.findall(src):
            # Normalise: on retire la partie paramétrée (ex: /admin/delete/<id> -> /admin/delete/)
            r_norm = re.sub(r"<[^>]+>", "", r)
            if not r_norm.endswith("/") and "/" in r and "<" in r:
                r_norm += "/"  # facilite les comparaisons par préfixe
            defined_routes.add(r_norm)

    # --- 4) ROUTES PARCOURSUP INDISPENSABLES ---
    required_parcoursup = [
        "/parcoursup",
        "/parcoursup/import",
        "/parcoursup/check",
        "/parcoursup/delete/",      # avec <cid>
        "/parcoursup/check-sms",
        "/parcoursup/logs/",        # avec <cid>
        "/brevo-sms-webhook",
        "/brevo-mail-webhook",
    ]
    missing_routes = []
    for need in required_parcoursup:
        ok = any(
            need == r or need.startswith(r) or r.startswith(need)
            for r in defined_routes
        )
        if not ok: missing_routes.append(need)
    if missing_routes:
        print("❌ Routes Parcoursup manquantes côté Flask :")
        for r in missing_routes: print("   •", r)
        sys.exit(1)
    print("✅ Routes Parcoursup détectées.")

    # --- 5) FETCH() DANS JS DOIVENT EXISTER EN ROUTES PYTHON ---
    js_fetch_targets = set()
    fetch_regex = re.compile(r"""fetch\(\s*['"](/[^'"]+)['"]""")
    for js_path in ["static/js/main_front.js", "static/js/main_admin.js"]:
        js = _read(js_path) or ""
        for u in fetch_regex.findall(js):
            # on ignore les querystrings et paramètres dynamiques
            u_clean = u.split("?")[0]
            js_fetch_targets.add(u_clean)

    # Construire une version "prefix-friendly" des routes
    def _matches_defined(u):
        # exemple: /admin/delete/123 doit matcher /admin/delete/
        for r in defined_routes:
            if u == r:
                return True
            # si la route définie est paramétrée, on a ajouté un trailing "/" plus haut
            if r.endswith("/") and u.startswith(r):
                return True
        return False

    bad = [u for u in js_fetch_targets if not _matches_defined(u)]
    if bad:
        print("❌ Des endpoints sont appelés en JS mais non définis en Python :")
        for b in sorted(bad): print("   •", b)
        sys.exit(1)
    print(f"✅ {len(js_fetch_targets)} endpoints JS → tous présents côté Flask.")

    # --- 6) TEMPLATE MAIL DE BASE DOIT AVOIR LES PLACEHOLDERS ---
    base_html_path = os.path.join(os.getcwd(), "templates", "email_base.html")
    base_mail = _read(base_html_path) or ""
    print(f"📂 Vérification email_base.html → {base_html_path}")

    # Autorise aussi {{ email_content | safe }}
    if ("{{ email_title" not in base_mail) or ("{{ email_content" not in base_mail):
        print("❌ templates/email_base.html doit contenir {{ email_title }} et {{ email_content }}.")
        sys.exit(1)
    if "logo_url" not in base_mail:
        print("⚠️  (recommandé) Utiliser {{ logo_url }} dans email_base.html pour l’image.")
    print("✅ email_base.html OK.")


    print("🎉 INTÉGRITÉ OK – Démarrage de l’application.\n")

# 👉 Appelle le contrôle immédiatement au démarrage :

# 🩵 Signal de vie immédiat pour Render
print("🚀 Lancement Flask en cours… (Render healthcheck OK)")
import sys; sys.stdout.flush()

run_startup_integrity_checks()
# ============================ FIN AUTO-CHECK ============================



load_dotenv()

# =====================================================
# 📎 GESTION DES PIÈCES JUSTIFICATIVES
# =====================================================

DOC_FIELDS = {
    "ci": ("fichiers_ci", "🪪 Carte d’identité / Passeport"),
    "photo": ("fichiers_photo", "📸 Photo d’identité"),
    "carte_vitale": ("fichiers_carte_vitale", "💳 Carte Vitale"),
    "cv": ("fichiers_cv", "📄 CV"),
    "lm": ("fichiers_lm", "🖋️ Lettre de motivation"),
}
# 📦 Préfixes utilisés pour nommer les fichiers de manière déterministe
FILE_PREFIX = {
    "ci": "CI",
    "photo": "Photo",
    "carte_vitale": "CarteVitale",
    "cv": "CV",
    "lm": "LettreMotivation",
}



def get_candidat(conn, cid):
    """Récupère un candidat complet par ID."""
    cur = conn.cursor()
    cur.execute("SELECT * FROM candidats WHERE id=?", (cid,))
    row = cur.fetchone()
    return dict(row) if row else None

def parse_list(v):
    """Convertit le JSON stocké des fichiers en vraie liste Python."""
    try:
        data = json.loads(v or "[]")
        if isinstance(data, list):
            return data
        return []
    except Exception:
        return []

def load_verif_docs(row):
    """Récupère le dictionnaire de vérification des pièces (conforme / non conforme)."""
    try:
        return json.loads(row.get("verif_docs") or "{}")
    except Exception:
        return {}


app = Flask(__name__)
# 🔒 Limite maximale pour les uploads (50 Mo)
app.config["MAX_CONTENT_LENGTH"] = 50 * 1024 * 1024

# 🛑 Erreur fichiers trop lourds (413)
@app.errorhandler(413)
def too_large(e):
    return (
        render_template(
            "error_upload.html",
            message="Le fichier que vous essayez d’envoyer est trop volumineux (maximum 8 Mo par fichier)."
        ),
        413
    )
     
app.secret_key = os.getenv("SECRET_KEY", "change-me")
# 🔐 Protection du module Parcoursup (toutes les routes /parcoursup)
@bp_parcoursup.before_request
def protect_parcoursup_routes():
    if not session.get("admin_ok"):
        return redirect(url_for("login"))

# 🔧 Rendre la fonction now() accessible dans les templates Jinja
@app.context_processor
def inject_now():
    return {'now': datetime.utcnow}

# 🕓 Filtre Jinja pour convertir une chaîne ISO en datetime
from datetime import datetime

@app.template_filter("to_datetime")
def to_datetime(value):
    if not value:
        return None
    try:
        return datetime.fromisoformat(value)
    except Exception:
        try:
            return datetime.strptime(value, "%Y-%m-%d %H:%M:%S")
        except Exception:
            return datetime.utcnow()

if "parcoursup" not in app.blueprints:
    app.register_blueprint(bp_parcoursup)
else:
    print("⚠️ Blueprint 'parcoursup' déjà enregistré — enregistrement ignoré.")


# =====================================================
# 🔤 Filtres Jinja pour affichage PDF
# =====================================================

@app.template_filter('modeemo')
def modeemo(value):
    if not value:
        return ""
    return "🏫 Présentiel" if value.lower() == "presentiel" else "💻 Distanciel"

@app.template_filter('btsfull')
def btsfull(value):
    if value == "MOS": return "BTS Management Opérationnel de la Sécurité (MOS)"
    if value == "MCO": return "BTS Management Commercial Opérationnel (MCO)"
    if value == "PI": return "BTS Professions Immobilières (PI)"
    if value == "CI": return "BTS Commerce International (CI)"
    if value == "NDRC": return "BTS Négociation et Digitalisation de la Relation Client (NDRC)"
    if value == "CG": return "BTS Comptabilité et Gestion (CG)"
    return value or ""

@app.template_filter('bacdisp')
def bacdisp(value):
    if not value: return ""
    if "prévu" in value.lower(): return "Prévu en 2026"
    if "oui" in value.lower(): return "Oui"
    if "non" in value.lower(): return "Non"
    return value

@app.template_filter('nirsp')
def nirsp(value):
    if not value: return ""
    value = value.replace(" ", "")
    return " ".join([value[i:i+2] for i in range(0, len(value), 2)]).strip()

@app.template_filter('dmy')
def dmy(value):
    try:
        from datetime import datetime
        return datetime.fromisoformat(value).strftime("%d/%m/%Y")
    except Exception:
        return value



# =====================================================
# 🎨 Filtres d'affichage Jinja unifiés
# =====================================================
import unicodedata

def _normalize(s: str) -> str:
    if not s: return ""
    return ''.join(c for c in unicodedata.normalize('NFD', s) if unicodedata.category(c) != 'Mn')

def _fmt_date_dmy(s: str) -> str:
    if not s: return ""
    try:
        if "T" in s or ":" in s:
            dt = datetime.fromisoformat(s.replace("Z",""))
            return dt.strftime("%d/%m/%Y")
        return datetime.strptime(s, "%Y-%m-%d").strftime("%d/%m/%Y")
    except Exception:
        return s

def _nir_spaced(nir: str) -> str:
    """Formate le NIR en groupes exacts ex: 1 93 09 74 256 233 79"""
    if not nir:
        return ""
    raw = (nir or "").strip().upper().replace(" ", "")
    groups, cuts, i = [], [1,2,2,2,3,3,2], 0
    for c in cuts:
        groups.append(raw[i:i+c])
        i += c
        if i >= len(raw):
            break
    return " ".join([g for g in groups if g])

def _bts_full(code: str) -> str:
    code = (code or "").strip().upper()
    return {
        "MCO": "BTS MANAGEMENT COMMERCIAL OPÉRATIONNEL (MCO)",
        "MOS": "BTS MANAGEMENT OPÉRATIONNEL DE LA SÉCURITÉ (MOS)",
        "PI":  "BTS PROFESSIONS IMMOBILIÈRES (PI)",
        "NDRC":"BTS NÉGOCIATION ET DIGITALISATION DE LA RELATION CLIENT (NDRC)",
        "CG":  "BTS COMPTABILITÉ ET GESTION (CG)",
        "CI":  "BTS COMMERCE INTERNATIONAL (CI)",
    }.get(code, code)

def _mode_emoji(mode: str) -> str:
    m = _normalize((mode or "").strip().lower())
    if "pres" in m: return "🏫 Présentiel"
    if "dist" in m: return "💻 Distanciel"
    return mode or ""

def _bac_status_display(v: str) -> str:
    s = _normalize((v or "").strip().lower())
    if s in {"oui","yes","true","1"}: return "Oui"
    if s in {"non","no","false","0"}: return "Non"
    if "prev" in s and "2026" in s: return "Prévu en 2026"
    if "prev" in s: return "Prévu"
    return (v or "").strip() or "—"

# 🔗 Enregistrement global
app.jinja_env.filters["dmy"]      = _fmt_date_dmy
app.jinja_env.filters["nirsp"]    = _nir_spaced
app.jinja_env.filters["btsfull"]  = _bts_full
app.jinja_env.filters["modeemo"]  = _mode_emoji
app.jinja_env.filters["bacdisp"]  = _bac_status_display



# =====================================================
# 💾 CONFIGURATION DU STOCKAGE PERSISTANT (Render)
# =====================================================

# 📁 Dossier persistant (base de données et PDF)
DATA_DIR = os.getenv("DATA_DIR", "/data")
os.makedirs(DATA_DIR, exist_ok=True)

# 📦 Base SQLite stockée sur le disque Render
DB_PATH = os.path.join(DATA_DIR, "app.db")

# 📂 Dossier des fichiers uploadés
# ⚠️ Pour qu’ils soient persistants aussi, on les place dans /data/uploads
UPLOAD_DIR = os.path.join(DATA_DIR, "uploads")
os.makedirs(UPLOAD_DIR, exist_ok=True)


ADMIN_PASSWORD = os.getenv("ADMIN_PASSWORD")  # optional

STATUTS = [
    ("preinscription", "Pré-inscription à traiter", "gray"),
    ("validee", "Candidature validée", "blue"),
    ("confirmee", "Inscription confirmée", "gold"),
    ("reconf_en_cours", "Reconfirmation en cours", "orange"),
    ("reconfirmee", "Inscription re-confirmée", "green"),
    ("annulee", "Inscription annulée", "red"),
    ("docs_non_conformes", "Docs non conformes", "black"),
]

def db():
    conn = sqlite3.connect(DB_PATH, timeout=10)
    conn.row_factory = sqlite3.Row
    return conn


def normalize_nir(nir: str) -> str:
    """Nettoie le NIR sans modifier sa structure (garde 2A/2B pour l'affichage)."""
    if not nir:
        return ""
    return nir.strip().upper().replace(" ", "")

def validate_nir(nir_raw, date_naissance_str, sexe_str):
    """Valide le NIR (clé + cohérence date + sexe). 
    Conserve 2A/2B dans la valeur stockée, mais effectue la vérification avec conversion temporaire."""
    if not nir_raw:
        return False, "Numéro de sécurité sociale manquant."

    nir_clean = normalize_nir(nir_raw)

    # 🔁 copie temporaire pour le calcul (convertit juste pour la vérif)
    nir_calc = nir_clean.replace("2A", "19").replace("2B", "18")
    digits = "".join(ch for ch in nir_calc if ch.isdigit())

    if len(digits) != 15:
        return False, "Le numéro de sécurité sociale doit comporter 15 caractères (chiffres ou A/B)."

    # ✅ Vérification de la clé de contrôle
    try:
        corps, cle = digits[:13], int(digits[13:15])
        calc = 97 - (int(corps) % 97)
        if calc != cle:
            return False, "Clé de contrôle du NIR incorrecte."
    except Exception:
        return False, "Format du NIR invalide."

    # ✅ Cohérence date
    if date_naissance_str:
        try:
            annee = int(date_naissance_str[:4])
            mois = int(date_naissance_str[5:7])
            yy_expected = str(annee)[-2:]
            mm_expected = f"{mois:02d}"
            yy, mm = digits[1:3], digits[3:5]

            if yy != yy_expected:
                return False, f"L'année ({yy}) ne correspond pas à la date ({yy_expected})."

            if mm.isdigit() and 1 <= int(mm) <= 12 and mm != mm_expected:
                return False, f"Le mois ({mm}) ne correspond pas au mois de naissance ({mm_expected})."
        except Exception:
            return False, "Date de naissance invalide."

    # ✅ Cohérence sexe
    try:
        s = int(digits[0])
        sexe_str = (sexe_str or "").lower()
        if (sexe_str.startswith("h") and s != 1) or (sexe_str.startswith("f") and s != 2):
            return False, "Le sexe indiqué ne correspond pas au NIR."
    except Exception:
        return False, "Format du NIR invalide."

    return True, ""




def init_db():
    conn = db()
    cur = conn.cursor()
    cur.execute("""
CREATE TABLE IF NOT EXISTS candidats (
    id TEXT PRIMARY KEY,
    numero_dossier TEXT,
    created_at TEXT,
    updated_at TEXT,
    nom TEXT, prenom TEXT, sexe TEXT,
    date_naissance TEXT, ville_naissance TEXT, cp_naissance TEXT, pays_naissance TEXT,
    num_secu TEXT, email TEXT, tel TEXT,
    adresse TEXT, cp TEXT, ville TEXT,
    bts TEXT, mode TEXT,
    bac_status TEXT, bac_type TEXT, bac_autre TEXT,
    permis_b INTEGER,
    est_mineur INTEGER,
    resp_nom TEXT, resp_prenom TEXT, resp_email TEXT, resp_tel TEXT,
    mos_parcours TEXT,
    aps_souhaitee INTEGER, aps_session TEXT,
    projet_pourquoi TEXT, projet_objectif TEXT, projet_passions TEXT,
    projet_qualites TEXT, projet_motivation TEXT, projet_recherche TEXT, projet_travail TEXT,
    fichiers_ci TEXT, fichiers_photo TEXT, fichiers_carte_vitale TEXT, fichiers_cv TEXT, fichiers_lm TEXT,
    statut TEXT,
    label_aps INTEGER, label_aut_ok INTEGER, label_cheque_ok INTEGER,
    token_confirm TEXT, token_confirm_exp TEXT,
    token_reconfirm TEXT, token_reconfirm_exp TEXT,
    commentaires TEXT,
    verif_docs TEXT, nouveau_doc INTEGER,
    replace_token TEXT, replace_token_exp TEXT, replace_meta TEXT,
    label_ypareo INTEGER, label_carte_etudiante INTEGER,
    date_validee TEXT, date_confirmee TEXT, date_reconfirmee TEXT,
    slug_public TEXT
);
""")

    cur.execute("""
    CREATE TABLE IF NOT EXISTS logs (
        id TEXT PRIMARY KEY,
        candidat_id TEXT,
        numero_dossier TEXT,
        type TEXT,
        payload TEXT,
        created_at TEXT
    );
    """)
    conn.commit()
    conn.close()

def ensure_schema():
    import sqlite3
    conn = sqlite3.connect(DB_PATH)
    cur = conn.cursor()

    # Vérifie les colonnes existantes
    cur.execute("PRAGMA table_info(candidats)")
    cols = [r[1] for r in cur.fetchall()]

    expected_cols = [
        "projet_qualites", "projet_motivation", "projet_recherche", "projet_travail",
        "verif_docs", "nouveau_doc", "replace_token", "replace_token_exp", "replace_meta",
        "label_ypareo", "label_carte_etudiante",
        "date_validee", "date_confirmee", "date_reconfirmee", "slug_public"
    ]

    for col in expected_cols:
        if col not in cols:
            print(f"🧱 Ajout automatique de la colonne {col} dans 'candidats'")
            cur.execute(f"ALTER TABLE candidats ADD COLUMN {col} TEXT DEFAULT ''")

    conn.commit()
    conn.close()

# ✅ Appel après définition
with app.app_context():
    init_db()
    ensure_schema()



import time
import sqlite3

# =====================================================
# 🧾 Vérifie et ajoute les colonnes apprentissage si manquantes
# =====================================================
def ensure_apprentissage_fields():
    conn = db()
    cur = conn.cursor()
    cur.execute("PRAGMA table_info(candidats)")
    cols = [r[1] for r in cur.fetchall()]

    if "entreprise_trouvee" not in cols:
        cur.execute("ALTER TABLE candidats ADD COLUMN entreprise_trouvee TEXT DEFAULT ''")
    if "recherches_commencees" not in cols:
        cur.execute("ALTER TABLE candidats ADD COLUMN recherches_commencees TEXT DEFAULT ''")

    conn.commit()
    conn.close()

# 🔁 Exécuter la vérification au démarrage
with app.app_context():
    ensure_apprentissage_fields()

# =====================================================
# 🎓 Vérifie et ajoute la colonne BACCALAURÉAT si manquante
# =====================================================
def ensure_baccalaureat_field():
    conn = db()
    cur = conn.cursor()
    cur.execute("PRAGMA table_info(candidats)")
    cols = [r[1] for r in cur.fetchall()]
    if "baccalaureat" not in cols:
        cur.execute("ALTER TABLE candidats ADD COLUMN baccalaureat TEXT DEFAULT ''")
        print("🧱 Colonne 'baccalaureat' ajoutée à la table 'candidats'")
    conn.commit()
    conn.close()

with app.app_context():
    ensure_baccalaureat_field()

# =====================================================
# 🧾 Vérifie et ajoute la colonne "cid" dans la table logs si manquante
# =====================================================
def ensure_logs_cid_column():
    conn = sqlite3.connect(DB_PATH)
    cur = conn.cursor()
    cur.execute("PRAGMA table_info(logs)")
    cols = [r[1] for r in cur.fetchall()]
    if "cid" not in cols:
        print("🧩 Ajout automatique de la colonne 'cid' dans logs…")
        cur.execute("ALTER TABLE logs ADD COLUMN cid TEXT")
        conn.commit()
    conn.close()

with app.app_context():
    ensure_logs_cid_column()


def log_event(candidat, type_, payload_dict):
    cid = candidat["id"] if isinstance(candidat, dict) else candidat
    nd = candidat.get("numero_dossier","") if isinstance(candidat, dict) else ""
    payload = json.dumps(payload_dict, ensure_ascii=False)
    attempt = 0
    while attempt < 5:
        try:
            conn = db()
            cur = conn.cursor()
            cur.execute(
                "INSERT INTO logs (id, candidat_id, numero_dossier, type, payload, created_at) VALUES (?, ?, ?, ?, ?, ?)",
                (str(uuid.uuid4()), cid, nd, type_, payload, datetime.now().isoformat())
            )
            conn.commit()
            conn.close()
            return
        except sqlite3.OperationalError as e:
            if "locked" in str(e).lower():
                attempt += 1
                print(f"⚠️ DB verrouillée (tentative {attempt}/5), nouvelle tentative dans 0.3s…")
                time.sleep(0.3)
            else:
                print("❌ Erreur log_event :", e)
                break
        except Exception as e:
            print("❌ Erreur inattendue log_event :", e)
            break


def get_counter_for_today(conn):
    cur = conn.cursor()
    cur.execute("SELECT COUNT(*) FROM candidats WHERE DATE(created_at) = DATE(?)", (datetime.now().isoformat(),))
    c = cur.fetchone()[0]
    return c + 1

def require_admin():
    return session.get("admin_ok", False)

# =====================================================
# 🕓 Ajout automatique de la colonne "last_relance" dans la table candidats
# =====================================================
def ensure_relance_field():
    conn = sqlite3.connect(DB_PATH)
    cur = conn.cursor()
    cur.execute("PRAGMA table_info(candidats)")
    cols = [r[1] for r in cur.fetchall()]
    if "last_relance" not in cols:
        print("🧩 Ajout automatique de la colonne 'last_relance' dans candidats…")
        cur.execute("ALTER TABLE candidats ADD COLUMN last_relance TEXT")
        conn.commit()
    conn.close()

with app.app_context():
    ensure_relance_field()



# =========================
# 🔐 PAGE DE CONNEXION SÉCURISÉE
# =========================
LOGIN_EMAIL = os.getenv("LOGIN_EMAIL", "clement@integraleacademy.com")
LOGIN_PASSWORD = os.getenv("LOGIN_PASSWORD", "Lv15052021@")


@app.route("/login", methods=["GET", "POST"])
def login():
    if request.method == "POST":
        email = request.form.get("email", "").strip().lower()
        password = request.form.get("password", "")
        if email == LOGIN_EMAIL and password == LOGIN_PASSWORD:
            session["admin_ok"] = True
            flash("Connexion réussie ✅", "success")
            return redirect(url_for("admin"))
        else:
            return render_template("login.html", error="Identifiants incorrects.")

    # si déjà connecté → admin
    if session.get("admin_ok"):
        return redirect(url_for("admin"))
    return render_template("login.html")


@app.route("/admin/login", methods=["GET", "POST"])
def admin_login():
    if not ADMIN_PASSWORD:
        return redirect(url_for("admin"))
    if request.method == "POST":
        pwd = request.form.get("password")
        if pwd == ADMIN_PASSWORD:
            session["admin_ok"] = True
            return redirect(url_for("admin"))
        flash("Mot de passe incorrect", "error")
    return render_template("admin_login.html", title="Admin – Connexion")

@app.route("/health")
def health():
    return "ok"

@app.route("/login-pole", methods=["GET", "POST"])
def pole_login():
    error = None

    if request.method == "POST":
        email = request.form.get("email", "").strip().lower()
        password = request.form.get("password", "").strip()

        # 🔐 Identifiants autorisés
        if email == "eric@polealternance.fr" and password == "polealternance":
            session["polealternance"] = True
            return redirect("/admin/pole-alternance")

        error = "Identifiants incorrects."

    return render_template("pole_login.html", error=error)


@app.route("/admin/pole-alternance")
def admin_pole_alternance():
    # 🔐 Protection d’accès
    if not session.get("polealternance") and session.get("email") != "clement@integraleacademy.com":
        return redirect("/login-pole")

    conn = db()
    cur = conn.cursor()

    # 🔍 On récupère uniquement les candidats accompagnés
    rows = cur.execute("""
        SELECT *
        FROM candidats
        WHERE souhaite_accompagnement = 'oui'
        ORDER BY created_at DESC
    """).fetchall()

    return render_template("admin_pole_alternance.html",
                           rows=rows,
                           statuts=STATUTS,
                           now=datetime.now)


@app.route("/admin/reprendre-plus-tard")
def admin_reprendre_plus_tard():
    # 🔐 Même protection que l’admin
    if not require_admin():
        return redirect("/login")

    DATA_DIR = os.getenv("DATA_DIR", "/data")
    DRAFT_PATH = os.path.join(DATA_DIR, "drafts.json")

    drafts = []
    if os.path.exists(DRAFT_PATH):
        try:
            with open(DRAFT_PATH, "r", encoding="utf-8") as f:
                drafts = json.load(f)
        except:
            drafts = []

    # 📌 Connexion DB pour vérifier les candidats existants
    conn = db()
    cur = conn.cursor()

    rows = []
    for d in drafts:
        form = d.get("full_form", {})

        email = form.get("email", "")

        # 🔍 Vérifie si l'email existe dans la table candidats (= inscrit)
        existing = cur.execute(
            "SELECT 1 FROM candidats WHERE email = ?", (email,)
        ).fetchone()

        deja_inscrit = 1 if existing else 0

        rows.append({
            "id": d.get("id"),
            "nom": form.get("nom", "—"),
            "prenom": form.get("prenom", "—"),
            "email": email,
            "tel": form.get("tel", "—"),
            "bts": form.get("bts", "—"),
            "updated_at": d.get("timestamp", "—"),
            "resume_link": d.get("resume_link", "#"),
            "deja_inscrit": deja_inscrit   # 👈 Pour afficher ✔️ / ❌
        })

    return render_template("admin_reprendre_plus_tard.html", rows=rows)



@app.route("/admin/reprendre-plus-tard/delete/<token>", methods=["POST"])
def admin_delete_draft(token):
    if not require_admin():
        abort(403)

    DATA_DIR = os.getenv("DATA_DIR", "/data")
    DRAFT_PATH = os.path.join(DATA_DIR, "drafts.json")

    drafts = []
    if os.path.exists(DRAFT_PATH):
        try:
            with open(DRAFT_PATH, "r", encoding="utf-8") as f:
                drafts = json.load(f)
        except:
            drafts = []

    # supprimer le bon token
    new_drafts = [d for d in drafts if d.get("id") != token]

    # écrire le fichier
    try:
        with open(DRAFT_PATH, "w", encoding="utf-8") as f:
            json.dump(new_drafts, f, indent=2, ensure_ascii=False)
    except Exception as e:
        print("❌ Erreur écriture drafts.json :", e)
        return jsonify({"ok": False, "error": str(e)}), 500

    print(f"🗑️ Brouillon supprimé : {token}")
    return jsonify({"ok": True})



# =====================================================
# 🎯 Vérifie et ajoute la colonne "souhaite_accompagnement" si manquante
# =====================================================
def ensure_souhaite_accompagnement_field():
    conn = db()
    cur = conn.cursor()
    cur.execute("PRAGMA table_info(candidats)")
    cols = [r[1] for r in cur.fetchall()]

    if "souhaite_accompagnement" not in cols:
        cur.execute("ALTER TABLE candidats ADD COLUMN souhaite_accompagnement TEXT DEFAULT ''")
        print("🧱 Colonne 'souhaite_accompagnement' ajoutée à la table 'candidats'")

    conn.commit()
    conn.close()


# ---------------- Public: Pré-inscriptions ----------------

@app.route("/")
def index():
    # =====================================================
    # 🧑‍💻 BYPASS ADMIN – permet d’accéder même si le portail est fermé
    # =====================================================
    if request.args.get("admin_bypass") == "1":
        print("🔓 Bypass admin activé – accès forcé au formulaire")
        return render_template(
            "index.html",
            title="Mode test (admin bypass)",
            portal_closed=False,
            portal_message="🧑‍💻 Formulaire visible uniquement pour test.",
            portal_comment="",
            saved_data=None,
            step=0
        )

    # =====================================================
    # 🔐 Vérifie l’état du portail
    # =====================================================
    portal = get_portal_status()
    if portal["status"] == "closed":
        return render_template(
            "index.html",
            title="Inscriptions momentanément fermées",
            portal_closed=True,
            portal_message=portal.get("message", ""),
            portal_comment=portal.get("comment", ""),
            saved_data=None,
            step=0
        )

    # =====================================================
    # 🚀 Portail ouvert → accès normal
    # =====================================================
    return render_template(
        "index.html",
        title="Pré-inscriptions BTS 2026",
        portal_closed=False,
        portal_message="",
        portal_comment="",
        saved_data=None,
        step=0
    )




def save_files(field_key: str, cand_id: str):
    """
    Sauvegarde les fichiers du champ `field_key` pour le candidat `cand_id`
    dans un dossier dédié : /uploads/<cand_id>/
    avec un nom clair : <PREFIX>_<cand_id>[_2].ext
    """
    files = request.files.getlist(field_key)
    saved = []
    prefix = FILE_PREFIX.get(field_key, field_key)

    # 📁 Crée un sous-dossier par candidat
    cand_dir = os.path.join(UPLOAD_DIR, cand_id)
    os.makedirs(cand_dir, exist_ok=True)

    idx = 1
    for f in files:
        if not f or not f.filename:
            continue
        _, ext = os.path.splitext(secure_filename(f.filename))
        base = f"{prefix}_{cand_id}{'' if idx == 1 else f'_{idx}'}{ext.lower()}"
        dest = os.path.join(cand_dir, base)
        f.save(dest)
        saved.append(dest)
        idx += 1

    return saved

# --- Convertit une valeur Oui/Non en booléen (1 ou 0) ---
def b(v):
    """Convertit les réponses 'Oui', 'True', '1', etc. en 1, sinon 0."""
    if v is None:
        return 0
    s = str(v).strip().lower()
    return 1 if s in ("on", "true", "1", "yes", "oui") else 0

# =====================================================
# 💾 SAUVEGARDE DU BROUILLON ("Reprendre plus tard")
# =====================================================
@app.route("/save_draft", methods=["POST"])
def save_draft():
    try:
        data = request.get_json()
        if not data:
            return jsonify({"success": False, "error": "Aucune donnée reçue"}), 400

        # 🔥 On sauvegarde TOUT, absolument tout :
        # tous les champs text, select, radios, checkboxes, etc.
        full_form = data.get("full_form", {})

        # 🔑 Id unique du brouillon (token)
        token = new_token()

        # 🔗 Lien de reprise
        resume_link = f"{request.url_root.rstrip('/')}/reprendre/{token}"

        # 🗂 Enregistrement complet du formulaire
        draft_data = {
            "id": token,
            "resume_link": resume_link,
            "full_form": full_form,
            "timestamp": datetime.now().isoformat()
        }

        DATA_FILE = os.path.join(DATA_DIR, "drafts.json")
        all_drafts = []

        if os.path.exists(DATA_FILE):
            try:
                with open(DATA_FILE, "r", encoding="utf-8") as f:
                    all_drafts = json.load(f)
            except:
                all_drafts = []

        all_drafts.append(draft_data)

        with open(DATA_FILE, "w", encoding="utf-8") as f:
            json.dump(all_drafts, f, indent=2, ensure_ascii=False)

        # 💌 Mail de reprise
        send_mail(
            full_form.get("email", ""),
            "Reprenez votre pré-inscription – Intégrale Academy",
            mail_html(
                "reprendre_plus_tard",
                prenom=full_form.get("prenom", ""),
                bts_label=full_form.get("bts", "BTS"),
                lien_espace=resume_link
            )
        )

        print(f"🟢 Brouillon complet enregistré — {resume_link}")
        return jsonify({"success": True, "link": resume_link})

    except Exception as e:
        print(f"❌ Erreur /save_draft : {e}")
        return jsonify({"success": False, "error": str(e)}), 500




# =====================================================
# 🔁 ROUTE DE REPRISE DU FORMULAIRE
# =====================================================
@app.route("/reprendre/<token>")
def reprendre_formulaire(token):
    DATA_DIR = os.getenv("DATA_DIR", "/data")
    DRAFT_PATH = os.path.join(DATA_DIR, "drafts.json")

    if not os.path.exists(DRAFT_PATH):
        abort(404)

    # 📂 Charger le fichier drafts.json
    try:
        with open(DRAFT_PATH, "r", encoding="utf-8") as f:
            drafts = json.load(f)
    except:
        drafts = []

    draft = next((d for d in drafts if d.get("id") == token), None)
    if not draft:
        abort(404)

    saved_full_form = draft.get("full_form", {})

    # 💥 On injecte TOUT dans la page
    return render_template(
        "index.html",
        saved_data=saved_full_form,   # <--- ici : les données
        step=1,
        portal_closed=False
    )


@app.route("/submit", methods=["POST"])
def submit():
    from datetime import date  # ✅ import ici (ou mets-le en haut du fichier)

    form = request.form
    conn = db()
    cur = conn.cursor()
    counter = get_counter_for_today(conn)
    numero = dossier_number(counter=counter)

    # === Normalisation du BAC ===
    baccalaureat = (form.get("baccalaureat") or "").strip()

    # === Projet motivé : nouveaux champs ===
    projet_pourquoi = (form.get("projet_pourquoi") or "").strip()
    projet_objectif = (form.get("projet_objectif") or "").strip()
    projet_passions = (form.get("projet_passions") or "").strip()

    # Cases à cocher
    qualites_list = request.form.getlist("qualites[]")
    motivation_list = request.form.getlist("motivation[]")
    valeurs_list = request.form.getlist("valeurs[]")
    travail_list = request.form.getlist("travail[]")

    projet_qualites = ", ".join(qualites_list)
    projet_motivation = ", ".join(motivation_list)
    projet_recherche = ", ".join(valeurs_list)
    projet_travail = ", ".join(travail_list)

    # === APS (sessions datées) ===
    aps_souhaitee = 1 if form.get("aps_souhaitee") == "oui" else 0
    aps_session_value = (form.get("aps_session") or "").strip()
    aps_session_other = (form.get("aps_session_other") or "").strip()
    raison_aps = (form.get("raison_aps") or "").strip()

    APS_SESSIONS = {
        "puget": "8 juillet → 12 août 2026 — Intégrale Academy (Puget-sur-Argens)",
        "autre": "7 septembre → 9 octobre 2026 — Intégrale Academy (Puget-sur-Argens)",
    }

    if aps_session_value in APS_SESSIONS:
        aps_session = APS_SESSIONS[aps_session_value]
    elif aps_session_value.lower() == "autre" and aps_session_other:
        aps_session = aps_session_other
    else:
        aps_session = aps_session_value

    entreprise_trouvee = (form.get("entreprise_trouvee") or "").strip()
    recherches_commencees = (form.get("recherches_commencees") or "").strip()
    souhaite_accompagnement = (form.get("souhaite_accompagnement") or "").strip()

    # 🎯 Vérification du numéro de sécurité sociale
    nir = form.get("num_secu", "").strip()
    date_naiss = form.get("date_naissance", "")
    sexe = form.get("sexe", "")

    # 🔞 Calcul automatique : mineur ou non
    est_mineur = 0
    try:
        naissance = datetime.strptime(date_naiss, "%Y-%m-%d").date()
        today = date.today()
        age = today.year - naissance.year - (
            (today.month, today.day) < (naissance.month, naissance.day)
        )
        if age < 18:
            est_mineur = 1
    except Exception:
        est_mineur = 0

    # 👤 Responsable légal (si mineur)
    resp_nom = (form.get("resp_nom") or "").strip()
    resp_prenom = (form.get("resp_prenom") or "").strip()
    resp_email = (form.get("resp_email") or "").strip()
    resp_tel = (form.get("resp_tel") or "").strip()

    # ✅ Maintenant seulement : form_overrides (car est_mineur/resp_* existent)
    form_overrides = {
        "projet_pourquoi": projet_pourquoi,
        "projet_objectif": projet_objectif,
        "projet_passions": projet_passions,
        "projet_qualites": projet_qualites,
        "projet_motivation": projet_motivation,
        "projet_recherche": projet_recherche,
        "projet_travail": projet_travail,

        # APS
        "aps_souhaitee": aps_souhaitee,
        "aps_session": aps_session,
        "raison_aps": raison_aps,
        "label_aps": aps_souhaitee,

        # Bac / apprentissage
        "bac_status": baccalaureat,
        "entreprise_trouvee": entreprise_trouvee,
        "recherches_commencees": recherches_commencees,
        "baccalaureat": baccalaureat,
        "souhaite_accompagnement": souhaite_accompagnement,

        # Mineur / responsable légal
        "est_mineur": est_mineur,
        "resp_nom": resp_nom,
        "resp_prenom": resp_prenom,
        "resp_email": resp_email,
        "resp_tel": resp_tel,
    }

    # 👍 Exception : ce numéro doit passer même s'il est invalide
    if nir == "123456789123456":
        print("⚠️ Bypass NIR activé pour un cas particulier (123456789123456)")
    else:
        ok, msg = validate_nir(nir, date_naiss, sexe)
        if not ok:
            flash(msg, "error")
            return redirect(request.referrer or url_for("index"))

    # 🧩 Création du candidat
    cand_id = str(uuid.uuid4())
    now = datetime.now().isoformat()

    # 📎 Sauvegarde des fichiers
    fichiers_ci = save_files("ci", cand_id)
    fichiers_photo = save_files("photo", cand_id)
    fichiers_carte_vitale = save_files("carte_vitale", cand_id)
    fichiers_cv = save_files("cv", cand_id)
    fichiers_lm = save_files("lm", cand_id)

    token_confirm = new_token()
    token_confirm_exp = (datetime.now() + timedelta(days=30)).isoformat()

    cur.execute("PRAGMA table_info(candidats)")
    cols = [r[1] for r in cur.fetchall()]
    placeholders = ",".join(["?"] * len(cols))
    sql = f"INSERT INTO candidats ({','.join(cols)}) VALUES ({placeholders})"

    values = []
    for c in cols:
        if c == "id":
            values.append(cand_id)
        elif c == "numero_dossier":
            values.append(numero)
        elif c in ("created_at", "updated_at"):
            values.append(now)
        elif "fichiers" in c:
            mapping = {
                "fichiers_ci": fichiers_ci,
                "fichiers_photo": fichiers_photo,
                "fichiers_carte_vitale": fichiers_carte_vitale,
                "fichiers_cv": fichiers_cv,
                "fichiers_lm": fichiers_lm,
            }
            values.append(json.dumps(mapping.get(c, [])))
        elif "verif_docs" in c or "replace_meta" in c:
            values.append("{}")
        elif c == "nouveau_doc":
            values.append(0)
        elif c in ("permis_b", "est_mineur"):
            values.append(int(form_overrides.get(c, 0)))
        elif "label" in c:
            values.append(form_overrides.get(c, 0))
        elif c == "statut":
            values.append("preinscription")
        elif c == "token_confirm":
            values.append(token_confirm)
        elif c == "token_confirm_exp":
            values.append(token_confirm_exp)
        elif c in ("token_reconfirm", "token_reconfirm_exp", "commentaires"):
            values.append("")
        elif c == "slug_public":
            values.append("")
        else:
            values.append(form_overrides.get(c, form.get(c, "")))

    cur.execute(sql, tuple(values))
    conn.commit()

    # =====================================================
    # 🤝 ENVOI AUTOMATIQUE – ACCOMPAGNEMENT PÔLE ALTERNANCE
    # =====================================================
    if souhaite_accompagnement.lower() == "oui":
        try:
            print("📤 Préparation du mail Pôle Alternance...")

            prenom = form.get("prenom", "")
            nom = form.get("nom", "")
            email = form.get("email", "")
            tel = form.get("tel", "")
            bts = form.get("bts", "")
            mode = form.get("mode", "")
            ville = form.get("ville", "")
            projet_pourquoi = form.get("projet_pourquoi", "")
            projet_objectif = form.get("projet_objectif", "")
            projet_passions = form.get("projet_passions", "")
            projet_qualites = ", ".join(request.form.getlist("qualites[]"))
            projet_motivation = ", ".join(request.form.getlist("motivation[]"))
            projet_recherche = ", ".join(request.form.getlist("valeurs[]"))
            projet_travail = ", ".join(request.form.getlist("travail[]"))

            bts_label = BTS_LABELS.get(bts.strip().upper(), bts)

            html = f"""
            <div style='font-family:Segoe UI,Arial,sans-serif;font-size:15px;color:#222;'>
              <h2 style='color:#2d2d2d;'>🎓 Nouveau candidat à accompagner Intégrale Academy</h2>
              <p>Bonjour,<br><br>
              Un nouveau candidat a demandé à être accompagné par Pôle Alternance pour trouver une entreprise (rentrée septembre 2026). Je vous prie de bien vouloir trouver ci-dessous les coordonnées du candidat. Vous trouverez également en pièce-jointe son CV et sa lettre de motivation. A bientôt, Clément VAILLANT (ceci est un mail automatique).</p>
              <ul>
                <li><strong>Nom :</strong> {nom}</li>
                <li><strong>Prénom :</strong> {prenom}</li>
                <li><strong>Email :</strong> {email}</li>
                <li><strong>Téléphone :</strong> {tel}</li>
                <li><strong>Ville :</strong> {ville}</li>
                <li><strong>BTS :</strong> {bts_label}</li>
                <li><strong>Mode :</strong> {mode}</li>
              </ul>
            </div>
            """

            attachments = []
            cand_dir = os.path.join(UPLOAD_DIR, cand_id)
            try:
                cv = next((f for f in os.listdir(cand_dir) if f.lower().startswith("cv_")), None)
                lm = next((f for f in os.listdir(cand_dir) if f.lower().startswith("lettremotivation_")), None)
                if cv: attachments.append(os.path.join(cand_dir, cv))
                if lm: attachments.append(os.path.join(cand_dir, lm))
            except Exception as e:
                print("⚠️ Impossible de lister les fichiers du candidat :", e)

            send_mail(
                "eric@polealternance.fr",
                f"🤝 Nouveau candidat à accompagner – {prenom} {nom}",
                html,
                attachments=attachments if attachments else None
            )

            print("✅ Mail Pôle Alternance envoyé avec succès.")
        except Exception as e:
            print("❌ Erreur envoi mail Pôle Alternance :", e)

    # 🧾 Logs et mails
    candidat = {
        "id": cand_id,
        "numero_dossier": numero,
        "email": form.get("email", ""),
        "prenom": form.get("prenom", "")
    }
    log_event(candidat, "PREINSCRIPTION_RECU", {"email": candidat["email"]})

    cur.execute("SELECT slug_public FROM candidats WHERE id=?", (cand_id,))
    row = cur.fetchone()
    slug = row[0] if row and row[0] else uuid.uuid4().hex[:10]
    if not row or not row[0]:
        cur.execute("UPDATE candidats SET slug_public=? WHERE id=?", (slug, cand_id))
        conn.commit()

    lien_espace = url_for("espace_candidat", slug=slug, _external=True)

    # 🔹 Préparation des libellés pour le mail
    bts_code = (form.get("bts") or "").strip().upper()
    bts_label = BTS_LABELS.get(bts_code, form.get("bts"))

    mode_raw = form.get("mode", "") or ""
    if "dist" in mode_raw.lower():
        mode_label_email = "💻 À distance 100% en ligne (visioconférence)"
    else:
        mode_label_email = "🏫 En présentiel à Puget-sur-Argens (Var, 83)"

    # ✉️ Mail accusé de réception
    html = mail_html(
        "accuse_reception",
        prenom=form.get("prenom", ""),
        bts_label=bts_label,
        lien_espace=lien_espace,

        numero_dossier=numero,
        form_nom=form.get("nom", ""),
        form_prenom=form.get("prenom", ""),
        form_email=form.get("email", ""),
        form_tel=form.get("tel", ""),
        form_mode_label=mode_label_email,
    )
    send_mail(form.get("email", ""), "Nous avons bien reçu votre pré-inscription – Intégrale Academy", html)

    # SMS accusé réception
    tel = (form.get("tel", "") or "").replace(" ", "")
    if tel.startswith("0"):
        tel = "+33" + tel[1:]
    bts_label = BTS_LABELS.get((form.get("bts") or "").strip().upper(), form.get("bts"))
    msg = sms_text(
        "accuse_reception",
        prenom=form.get("prenom", ""),
        bts_label=bts_label,
        lien_espace=lien_espace
    )
    send_sms_brevo(tel, msg)
    log_event(candidat, "SMS_ENVOYE", {"type": "accuse_reception", "tel": tel})
    log_event(candidat, "MAIL_ENVOYE", {"type": "accuse_reception"})

    # Notification email to the admin inbox disabled on request.

    return redirect(lien_espace)








# ---------------- Admin ----------------

@app.route("/admin")
def admin():
    if not require_admin():
        return redirect(url_for("login"))

    q = request.args.get("q","").strip().lower()
    flt_bts = request.args.get("bts","")
    flt_statut = request.args.get("statut","")
    flt_mode = request.args.get("mode","")
    flt_label = request.args.get("label","")
    flt_relance = request.args.get("relances","")

    conn = db()
    cur = conn.cursor()
    cur.execute("SELECT * FROM candidats ORDER BY created_at DESC")
    rows = [dict(r) for r in cur.fetchall()]

    def match(row):
        ok = True
        if q:
            blob = " ".join([str(row.get(k,'')) for k in ("nom","prenom","email","tel","bts","numero_dossier")]).lower()
            ok &= q in blob
        if flt_bts:
            ok &= row.get("bts","") == flt_bts
        if flt_statut:
            ok &= row.get("statut","") == flt_statut
        if flt_mode:
            ok &= row.get("mode","") == flt_mode
        if flt_relance:
            ok &= row.get("last_relance") not in (None, "", "null")
        if flt_label == "APS":
            ok &= row.get("label_aps",0) == 1
        if flt_label == "AUT":
            ok &= row.get("label_aut_ok",0) == 1
        if flt_label == "CHEQUE":
            ok &= row.get("label_cheque_ok",0) == 1
        return ok

    rows = [r for r in rows if match(r)]
    return render_template("admin.html", title="Administration", rows=rows, statuts=STATUTS)

@app.get("/api/kpi")
def api_kpi():
    """
    Renvoie les stats pour la barre de KPI en haut de l'admin.
    Les clés renvoyées correspondent aux statuts dans la table `candidats`.
    """

    import sqlite3  # au cas où ça ne soit pas déjà importé tout en haut

    # ⚠️ Adapte DB_PATH si ton projet utilise un autre nom de variable
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    cur = conn.cursor()

    # 🧮 Statuts à compter
    stats = {
        "preinscription": 0,
        "validee": 0,
        "confirmee": 0,
        "reconf_en_cours": 0,
        "reconfirmee": 0,
        "annulee": 0,
        "docs_non_conformes": 0,
    }

    for statut in stats.keys():
        cur.execute("SELECT COUNT(*) AS c FROM candidats WHERE statut = ?", (statut,))
        row = cur.fetchone()
        stats[statut] = row["c"] if row else 0

    conn.close()
    return jsonify(stats)

# =====================================================
# 🔁 API JSON – Liste simplifiée des candidats pour refresh auto
# =====================================================
@app.route("/admin/json")
def admin_json():
    if not require_admin():
        abort(403)
    conn = db()
    rows = [dict(r) for r in conn.execute("SELECT id, nom, prenom, bts, mode, tel, email, statut, created_at FROM candidats ORDER BY created_at DESC LIMIT 500")]
    conn.close()
    return jsonify({"ok": True, "rows": rows})

# =====================================================
# ⚡ API JSON – Détail d'une ligne candidat (refresh dynamique)
# =====================================================
@app.route("/admin/row/<cid>")
def admin_row(cid):
    if not require_admin():
        abort(403)
    conn = db()
    row = conn.execute("SELECT * FROM candidats WHERE id=?", (cid,)).fetchone()
    conn.close()
    if not row:
        return jsonify(ok=False, error="Candidat introuvable"), 404

    row = dict(row)

    # ✅ Indique au front si le candidat a la carte étudiante cochée
    row["has_badge_carte"] = bool(
    row.get("label_carte_etudiante")
    or row.get("label_carte_ok")
    or row.get("carte_etudiante")
)

    # 🔹 Liste des statuts disponibles pour le <select>
    statuts = [{"key": s[0], "label": s[1]} for s in STATUTS]

    return jsonify(ok=True, row=row, statuts=statuts)


@app.route("/admin/update-field", methods=["POST"])
def admin_update_field():
    if not require_admin():
        abort(403)

    data = request.get_json(force=True)
    cid = data.get("id")
    field = data.get("field")
    value = data.get("value")

    allowed = [
        "nom", "prenom", "bts", "mode", "tel", "email",
        "label_aps", "label_aut_ok", "label_cheque_ok",
        "label_ypareo", "label_carte_etudiante",
        "commentaires", "nouveau_doc"
    ]

    if field not in allowed:
        return jsonify({"ok": False, "error": "Champ non autorisé"}), 400

    try:
        # 🧩 On garde le type exact selon le champ
        if field.startswith("label_"):
            # c’est une case à cocher → booléen
            value = 1 if str(value).lower() in ("true", "1", "on", "yes") else 0
        else:
            # c’est un champ texte
            value = str(value).strip()

        conn = db()
        cur = conn.cursor()
        cur.execute(f"UPDATE candidats SET {field}=?, updated_at=? WHERE id=?",
                    (value, datetime.now().isoformat(), cid))
        conn.commit()

        print(f"🟢 UPDATE-FIELD: id={cid}, field={field}, value={value}")

        # Log interne
        cur.execute("SELECT * FROM candidats WHERE id=?", (cid,))
        row = dict(cur.fetchone())
        log_event(row, "FIELD_UPDATE", {"field": field, "value": value})

        conn.close()
        return jsonify({"ok": True, "value": value})

    except Exception as e:
        print("❌ Erreur update-field :", e)
        return jsonify({"ok": False, "error": str(e)}), 500

@app.route("/admin/update-status", methods=["POST"])
def admin_update_status():
    if not require_admin():
        abort(403)

    data = request.json or {}
    cid = data.get("id")
    value = data.get("value")

    conn = db()
    cur = conn.cursor()

    # 🔽 On récupère l'ancien statut AVANT toute modification
    cur.execute("SELECT * FROM candidats WHERE id=?", (cid,))
    full_row = dict(cur.fetchone())

    # 🛑 ANTI-DOUBLON : si le statut est déjà celui demandé → NE RIEN ENVOYER
    if full_row.get("statut") == value:
        return jsonify({"ok": True, "statut": value})

    # 🕓 Enregistre la date correspondant au statut
    now_iso = datetime.now().isoformat()
    now_str = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

    if value == "validee":
        cur.execute(
            "UPDATE candidats SET statut=?, date_validee=?, updated_at=? WHERE id=?",
            (value, now_str, now_iso, cid)
        )
    elif value == "confirmee":
        cur.execute(
            "UPDATE candidats SET statut=?, date_confirmee=?, updated_at=? WHERE id=?",
            (value, now_str, now_iso, cid)
        )
    elif value == "reconfirmee":
        cur.execute(
            "UPDATE candidats SET statut=?, date_reconfirmee=?, updated_at=? WHERE id=?",
            (value, now_str, now_iso, cid)
        )
    else:
        cur.execute(
            "UPDATE candidats SET statut=?, updated_at=? WHERE id=?",
            (value, now_iso, cid)
        )

    conn.commit()

    # 🧹 Suppression badge relance
    if value in ["confirmee", "reconfirmee", "validee"]:
        try:
            cur.execute(
                "UPDATE candidats SET last_relance=NULL, updated_at=? WHERE id=?",
                (datetime.now().isoformat(), cid)
            )
            conn.commit()
        except Exception as e:
            print("⚠️ Erreur suppression badge relance :", e)

    # 🔄 Recharge données fraîches
    cur.execute(
        "SELECT statut, date_validee, date_confirmee, date_reconfirmee, last_relance FROM candidats WHERE id=?",
        (cid,)
    )
    row = dict(cur.fetchone())

    # ✉️📱 ENVOIS AUTOMATIQUES SELON STATUT
    # IMPORTANT : maintenant qu’on est sûr que ce n'est pas un doublon → on envoie

    # 1️⃣ CANDIDATURE VALIDÉE
    if value == "validee":
        token = full_row.get("token_confirm")
        if not token:
            token = new_token()
            exp = (datetime.now() + timedelta(days=30)).isoformat()
            cur.execute(
                "UPDATE candidats SET token_confirm=?, token_confirm_exp=?, updated_at=? WHERE id=?",
                (token, exp, datetime.now().isoformat(), cid)
            )
            conn.commit()
            cur.execute("SELECT * FROM candidats WHERE id=?", (cid,))
            full_row = dict(cur.fetchone())

        lien_confirmation = make_signed_link("/confirm-inscription", token)

        BASE_URL = os.getenv("BASE_URL", "https://inscriptionsbts.onrender.com").rstrip("/")
        slug = full_row.get("slug_public")
        lien_espace = f"{BASE_URL}/espace/{slug}"

        ctx = get_mail_context(full_row, lien_espace=lien_espace, lien_confirmation=lien_confirmation)
        ctx["bts_label"] = BTS_LABELS.get(ctx["bts_label"], ctx["bts_label"])

        html = mail_html("candidature_validee", **ctx)
        send_mail(full_row["email"], "Votre candidature est validée – Confirmez votre inscription", html)

        tel = (full_row.get("tel", "") or "").replace(" ", "")
        if tel.startswith("0"):
            tel = "+33" + tel[1:]
        sms_msg = sms_text("candidature_validee", prenom=ctx["prenom"], bts_label=ctx["bts_label"],
                           lien_espace=lien_espace, lien_confirmation=lien_confirmation)
        send_sms_brevo(tel, sms_msg)

    # 2️⃣ INSCRIPTION CONFIRMÉE
    elif value == "confirmee":
        cur.execute("SELECT * FROM candidats WHERE id=?", (cid,))
        full_row = dict(cur.fetchone())

        ctx = get_mail_context(full_row)
        ctx["bts_label"] = BTS_LABELS.get(ctx["bts_label"], ctx["bts_label"])

        html = mail_html("inscription_confirmee", **ctx)
        send_mail(full_row["email"], "Inscription confirmée – Intégrale Academy", html)

        tel = (full_row.get("tel", "") or "").replace(" ", "")
        if tel.startswith("0"):
            tel = "+33" + tel[1:]
        sms_msg = sms_text("inscription_confirmee", prenom=ctx["prenom"], bts_label=ctx["bts_label"])
        send_sms_brevo(tel, sms_msg)

    # 3️⃣ RECONFIRMATION
    elif value == "reconfirmee":
        ctx = get_mail_context(full_row)
        ctx["bts_label"] = BTS_LABELS.get(ctx["bts_label"], ctx["bts_label"])

        html = render_template("mail_bienvenue.html", prenom=ctx["prenom"], bts=full_row["bts"])
        send_mail(full_row["email"], "Bienvenue à Intégrale Academy 🎓", html)

    # LOG STATUT
    log_event(full_row, "STATUT_CHANGE", {"statut": value})

    conn.close()

    # Réponse finale
    return jsonify({
        "ok": True,
        **row
    })






    # =====================================================
    # 📩 ENVOIS AUTOMATIQUES SELON STATUT
    # =====================================================

    # 1️⃣ CANDIDATURE VALIDÉE
    if value == "validee":

        # 🔐 Token de confirmation
        token = full_row.get("token_confirm")
        if not token:
            token = new_token()
            exp = (datetime.now() + timedelta(days=30)).isoformat()
            cur.execute(
                "UPDATE candidats SET token_confirm=?, token_confirm_exp=?, updated_at=? WHERE id=?",
                (token, exp, datetime.now().isoformat(), cid)
            )
            conn.commit()
            cur.execute("SELECT * FROM candidats WHERE id=?", (cid,))
            full_row = dict(cur.fetchone())

        # 👉 Lien confirmation
        lien_confirmation = make_signed_link("/confirm-inscription", token)

        # 👉 Lien espace candidat
        BASE_URL = os.getenv("BASE_URL", "https://inscriptionsbts.onrender.com").rstrip("/")
        slug = full_row.get("slug_public")
        lien_espace = f"{BASE_URL}/espace/{slug}"

        # 🎁 Contexte universel
        ctx = get_mail_context(full_row, lien_espace=lien_espace, lien_confirmation=lien_confirmation)
        ctx["bts_label"] = BTS_LABELS.get(ctx["bts_label"], ctx["bts_label"])

        # ✉ MAIL
        html = mail_html("candidature_validee", **ctx)
        send_mail(full_row["email"], "Votre candidature est validée – Confirmez votre inscription", html)

        # 📱 SMS
        tel = (full_row.get("tel", "") or "").replace(" ", "")
        if tel.startswith("0"):
            tel = "+33" + tel[1:]
        sms_msg = sms_text(
            "candidature_validee",
            prenom=ctx["prenom"],
            bts_label=ctx["bts_label"],
            lien_espace=lien_espace,
            lien_confirmation=lien_confirmation,
        )
        send_sms_brevo(tel, sms_msg)


    # 2️⃣ INSCRIPTION CONFIRMÉE
    elif value == "confirmee":

        # 🔄 Recharge les données fraîches
        cur.execute("SELECT * FROM candidats WHERE id=?", (cid,))
        full_row = dict(cur.fetchone())

        # 🎁 Contexte complet
        ctx = get_mail_context(full_row)
        ctx["bts_label"] = BTS_LABELS.get(ctx["bts_label"], ctx["bts_label"])

        # ✉ MAIL confirmation
        html = mail_html("inscription_confirmee", **ctx)
        send_mail(full_row["email"], "Inscription confirmée – Intégrale Academy", html)

        # 📱 SMS
        tel = (full_row.get("tel", "") or "").replace(" ", "")
        if tel.startswith("0"):
            tel = "+33" + tel[1:]
        sms_msg = sms_text(
            "inscription_confirmee",
            prenom=ctx["prenom"],
            bts_label=ctx["bts_label"],
        )
        send_sms_brevo(tel, sms_msg)



    # 3️⃣ RECONFIRMATION VALIDÉE
    elif value == "reconfirmee":

        ctx = get_mail_context(full_row)
        ctx["bts_label"] = BTS_LABELS.get(ctx["bts_label"], ctx["bts_label"])

        merci_html = render_template("mail_bienvenue.html", prenom=ctx["prenom"], bts=full_row["bts"])
        send_mail(full_row["email"], "Bienvenue à Intégrale Academy 🎓", merci_html)


    # LOG GENERAL
    log_event(full_row, "STATUT_CHANGE", {"statut": value})

    conn.close()

    # 🔁 Réponse front
    return jsonify({
        "ok": True,
        "statut": row.get("statut"),
        "date_validee": row.get("date_validee"),
        "date_confirmee": row.get("date_confirmee"),
        "date_reconfirmee": row.get("date_reconfirmee"),
        "last_relance": row.get("last_relance"),
    })






@app.route("/admin/print/<cid>")
def admin_print(cid):
    if not require_admin(): abort(403)
    from weasyprint import HTML
    conn = db(); cur = conn.cursor()
    cur.execute("SELECT * FROM candidats WHERE id=?", (cid,))
    row = dict(cur.fetchone())
    html = render_template("pdf_fiche.html", row=row, now=datetime.now())
    pdf = HTML(string=html, base_url=url_for('index', _external=True)).write_pdf()
    fname = f"fiche_{row.get('numero_dossier','')}.pdf"
    path = os.path.join(DATA_DIR, fname)
    with open(path, "wb") as f:
        f.write(pdf)
    log_event(row, "PDF_GENERE", {"file": fname})
    return send_file(path, mimetype="application/pdf", as_attachment=True, download_name=fname)

@app.route("/admin/delete/<cid>", methods=["POST"])
def admin_delete(cid):
    if not require_admin(): abort(403)
    conn = db(); cur = conn.cursor()
    cur.execute("SELECT * FROM candidats WHERE id=?", (cid,))
    row = cur.fetchone()
    if row:
        row = dict(row)
        cur.execute("DELETE FROM candidats WHERE id=?", (cid,))
        conn.commit()
        log_event(row, "FICHE_SUPPRIMEE", {})
    return jsonify({"ok":True})

@app.route("/admin/export.csv")
def admin_export_csv():
    if not require_admin(): abort(403)
    import csv, io
    conn = db(); cur = conn.cursor()
    cur.execute("SELECT * FROM candidats ORDER BY created_at DESC")
    rows = [dict(r) for r in cur.fetchall()]
    buf = io.StringIO()
    if rows:
        w = csv.DictWriter(buf, fieldnames=list(rows[0].keys()))
        w.writeheader()
        for r in rows:
            w.writerow(r)
    return buf.getvalue(), 200, {"Content-Type":"text/csv; charset=utf-8", "Content-Disposition":"attachment; filename=export.csv"}

@app.route("/admin/export.json")
def admin_export_json():
    if not require_admin(): abort(403)
    conn = db(); cur = conn.cursor()
    cur.execute("SELECT * FROM candidats ORDER BY created_at DESC")
    rows = [dict(r) for r in cur.fetchall()]
    return jsonify(rows)

# =====================================================
# 📘 NOMS COMPLETS DES BTS
# =====================================================
BTS_LABELS = {
    "MCO": "BTS MANAGEMENT COMMERCIAL OPÉRATIONNEL (MCO)",
    "MOS": "BTS MANAGEMENT OPÉRATIONNEL DE LA SÉCURITÉ (MOS)",
    "PI": "BTS PROFESSIONS IMMOBILIÈRES (PI)",
    "NDRC": "BTS NÉGOCIATION ET DIGITALISATION DE LA RELATION CLIENT (NDRC)",
    "CG": "BTS COMPTABILITÉ ET GESTION (CG)",
    "CI": "BTS COMMERCE INTERNATIONAL (CI)"
}


# =====================================================
# 🧾 GÉNÉRATION CERTIFICAT DE SCOLARITÉ (DOCX UNIQUEMENT)
# =====================================================
@app.route("/admin/generate_certificat/<id>")
def admin_generate_certificat(id):
    from docx import Document
    from datetime import datetime
    from flask import send_file

    # 📂 chemins
    template_path = os.path.join("static", "templates", "certificat de scolarité 2026.docx")
    output_dir = os.path.join(DATA_DIR, "certificats")
    os.makedirs(output_dir, exist_ok=True)

    # 🧾 récupérer infos candidat
    conn = db()
    cur = conn.cursor()
    cur.execute("SELECT nom, prenom, bts FROM candidats WHERE id = ?", (id,))
    row = cur.fetchone()
    conn.close()
    if not row:
        return "Candidat introuvable", 404

    nom, prenom, bts = row
    full_name = f"{prenom.upper()} {nom.upper()}"
    date_now = datetime.now().strftime("%d/%m/%Y")

    # 🧩 Nom complet du BTS
    bts_nom_complet = BTS_LABELS.get(bts.strip().upper(), bts)

    # 🔧 valeurs de remplacement
    replacements = {
        "{{NOM_PRENOM}}": full_name,
        "{{FORMATION}}": bts_nom_complet,
        "{{DATE_AUJOURDHUI}}": date_now,
        "{{ANNEE_DEBUT}}": "2026",
        "{{ANNEE_FIN}}": "2028",
    }

    # 🧩 ouvrir modèle Word et remplacer les balises dans tous les paragraphes et tables
    doc = Document(template_path)

    def replace_text_in_paragraph(paragraph):
        for key, val in replacements.items():
            if key in paragraph.text:
                for run in paragraph.runs:
                    if key in run.text:
                        run.text = run.text.replace(key, val)

    for p in doc.paragraphs:
        replace_text_in_paragraph(p)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_text_in_paragraph(p)

    # 💾 sauvegarde du document Word généré
    output_docx = os.path.join(output_dir, f"certificat_{id}.docx")
    doc.save(output_docx)

    # 🧩 Log d’action
    log_event({"id": id}, "DOC_GENERE", {
        "type": "certificat_scolarite_distanciel",
        "file": output_docx
    })

    print(f"✅ Certificat généré pour {full_name}")
    return send_file(output_docx, as_attachment=True)


# =====================================================
# 🏫 CERTIFICAT DE SCOLARITÉ PRÉSENTIEL
# =====================================================
@app.route("/admin/generate_certificat_presentiel/<id>")
def admin_generate_certificat_presentiel(id):
    from docx import Document
    from datetime import datetime
    from flask import send_file

    # 📂 chemins
    template_path = os.path.join("static", "templates", "certificat_scolarite_presentiel.docx")
    output_dir = os.path.join(DATA_DIR, "certificats")
    os.makedirs(output_dir, exist_ok=True)

    # 🧾 récupérer infos candidat
    conn = db()
    cur = conn.cursor()
    cur.execute("SELECT nom, prenom, bts FROM candidats WHERE id = ?", (id,))
    row = cur.fetchone()
    conn.close()
    if not row:
        return "Candidat introuvable", 404

    nom, prenom, bts = row
    full_name = f"{prenom.upper()} {nom.upper()}"
    date_now = datetime.now().strftime("%d/%m/%Y")

    # 🧩 Nom complet du BTS
    bts_nom_complet = BTS_LABELS.get(bts.strip().upper(), bts)

    # 🔧 valeurs de remplacement
    replacements = {
        "{{NOM_PRENOM}}": full_name,
        "{{FORMATION}}": bts_nom_complet,
        "{{DATE_AUJOURDHUI}}": date_now,
        "{{ANNEE_DEBUT}}": "2026",
        "{{ANNEE_FIN}}": "2028",
    }

    # 🧩 ouvrir modèle Word et remplacer les balises dans tous les paragraphes et tables
    doc = Document(template_path)

    def replace_text_in_paragraph(paragraph):
        for key, val in replacements.items():
            if key in paragraph.text:
                for run in paragraph.runs:
                    if key in run.text:
                        run.text = run.text.replace(key, val)

    for p in doc.paragraphs:
        replace_text_in_paragraph(p)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_text_in_paragraph(p)

    # 💾 sauvegarde du document Word généré
    output_docx = os.path.join(output_dir, f"certificat_presentiel_{id}.docx")
    doc.save(output_docx)

    # 🧩 Log d’action
    log_event({"id": id}, "DOC_GENERE", {
        "type": "certificat_scolarite_presentiel",
        "file": output_docx
    })

    print(f"✅ Certificat présentiel généré pour {full_name}")
    return send_file(output_docx, as_attachment=True)





# =====================================================
# ✉️ ENVOI DU CERTIFICAT DE SCOLARITÉ PAR MAIL
# =====================================================
@app.route("/admin/send_certificat/<id>")
def admin_send_certificat(id):
    from flask import jsonify
    from utils import send_mail
    import os

    # 📂 Chemins des certificats
    cert_dir = os.path.join(DATA_DIR, "certificats")
    cert_path = os.path.join(cert_dir, f"certificat_{id}.docx")

    # 🔍 Vérifier que le fichier existe
    if not os.path.exists(cert_path):
        return jsonify({"ok": False, "error": "Le certificat n’a pas encore été généré."}), 404

    # 🧾 Récupérer les infos du candidat
    conn = db()
    cur = conn.cursor()
    cur.execute("SELECT prenom, nom, email, bts FROM candidats WHERE id = ?", (id,))
    row = cur.fetchone()
    conn.close()

    if not row:
        return jsonify({"ok": False, "error": "Candidat introuvable"}), 404

    prenom, nom, email, bts = row
    bts_nom_complet = BTS_LABELS.get(bts.strip().upper(), bts)

    # ✉️ Mail HTML
    html = mail_html(
        "certificat",
        prenom=prenom.title(),
        bts_label=bts_nom_complet
    )
    subject = f"Votre certificat de scolarité – {bts_nom_complet} 2026-2028"

    try:
        send_mail(email, subject, html, attachments=[cert_path])
        log_event({"id": id}, "MAIL_ENVOYE", {"type": "certificat"})
        print(f"✅ Certificat envoyé à {prenom} {nom} ({email})")
        return jsonify({"ok": True})
    except Exception as e:
        print(f"❌ Erreur envoi certificat à {prenom} {nom} :", e)
        return jsonify({"ok": False, "error": str(e)}), 500


# =====================================================
# ✉️ ENVOI DU CERTIFICAT DE SCOLARITÉ PRÉSENTIEL PAR MAIL
# =====================================================
@app.route("/admin/send_certificat_presentiel/<id>")
def admin_send_certificat_presentiel(id):
    from flask import jsonify
    from utils import send_mail
    import os

    cert_dir = os.path.join(DATA_DIR, "certificats")
    cert_path = os.path.join(cert_dir, f"certificat_presentiel_{id}.docx")

    if not os.path.exists(cert_path):
        return jsonify({"ok": False, "error": "Le certificat présentiel n’a pas encore été généré."}), 404

    conn = db()
    cur = conn.cursor()
    cur.execute("SELECT prenom, nom, email, bts FROM candidats WHERE id = ?", (id,))
    row = cur.fetchone()
    conn.close()

    if not row:
        return jsonify({"ok": False, "error": "Candidat introuvable"}), 404

    prenom, nom, email, bts = row
    bts_nom_complet = BTS_LABELS.get(bts.strip().upper(), bts)

    html = mail_html(
        "certificat_presentiel",
        prenom=prenom.title(),
        bts_label=bts_nom_complet
    )
    subject = f"Votre certificat de scolarité – Présentiel ({bts_nom_complet} 2026-2028)"

    try:
        send_mail(email, subject, html, attachments=[cert_path])
        log_event({"id": id}, "MAIL_ENVOYE", {"type": "certificat_presentiel"})
        print(f"✅ Certificat présentiel envoyé à {prenom} {nom} ({email})")
        return jsonify({"ok": True})
    except Exception as e:
        print(f"❌ Erreur envoi certificat présentiel à {prenom} {nom} :", e)
        return jsonify({"ok": False, "error": str(e)}), 500


# =====================================================
# 🕓 HISTORIQUE DES ACTIONS (LOGS)
# =====================================================
@app.route("/admin/logs/<cid>")
def admin_logs(cid):
    if not require_admin():
        abort(403)

    conn = db()
    cur = conn.cursor()
    cur.execute("""
        SELECT type, payload, created_at 
        FROM logs 
        WHERE candidat_id=? 
        ORDER BY datetime(created_at) DESC
    """, (cid,))
    rows = [dict(r) for r in cur.fetchall()]
    conn.close()

    # 🧩 Nettoyage : on affiche un résumé lisible
    for r in rows:
        try:
            p = json.loads(r["payload"])
            if isinstance(p, dict):
                r["payload"] = " / ".join(f"{k}: {v}" for k, v in p.items())
        except Exception:
            pass
    return jsonify(rows)


# =====================================================
# 📎 GESTION SIMPLIFIÉE DES PIÈCES JUSTIFICATIVES
# =====================================================

@app.route("/admin/files/<cid>")
def admin_files(cid):
    if not require_admin():
        abort(403)

    conn = db()
    row = get_candidat(conn, cid)
    if not row:
        conn.close()
        abort(404)

    verif = load_verif_docs(row)
    files_data = []

    # === Étape 1 : Fichiers enregistrés en base ===
    for key, (field, label) in DOC_FIELDS.items():
        file_list = parse_list(row.get(field))
        for path in file_list:
            fname = os.path.basename(path)
            status_info = verif.get(fname, {})
            files_data.append({
                "type": key,
                "label": label,
                "filename": fname,
                "path": path,
                "status": status_info.get("etat", "en_attente"),
                "horodatage": status_info.get("horodatage", "")
            })

    # === Étape 2 : Détection automatique des nouveaux fichiers dans le dossier du candidat ===
    cand_dir = os.path.join(UPLOAD_DIR, cid)
    try:
        all_on_disk = os.listdir(cand_dir)
    except Exception as e:
        print(f"⚠️ Erreur lecture du dossier candidat {cid} : {e}")
        all_on_disk = []

    existing_filenames = {os.path.basename(f["filename"]) for f in files_data}
    nouveaux_detectes = []

    for f in all_on_disk:
        if f not in existing_filenames:
            full_path = os.path.join(cand_dir, f)
            if os.path.isfile(full_path):
                files_data.append({
                    "type": "nouveau",
                    "label": "📥 Nouveau document déposé",
                    "filename": f,
                    "path": full_path,
                    "status": "nouveau",
                    "horodatage": datetime.fromtimestamp(os.path.getmtime(full_path)).strftime("%d/%m/%Y à %H:%M")
                })
                nouveaux_detectes.append(f)

    # === Étape 3 : Marquage des fichiers présents dans replace_meta (si applicable) ===
    if row.get("replace_meta"):
        try:
            meta = json.loads(row["replace_meta"])
            nouveaux_fichiers = [x["fichier"] for x in meta.get("nouveaux_fichiers", [])]
            for f in files_data:
                if f["filename"] in nouveaux_fichiers:
                    f["type"] = "nouveau"
                    f["label"] = f"📥 Nouveau document déposé — {f['label']}"
        except Exception as e:
            print("⚠️ Erreur détection replace_meta :", e)

    # === Étape 4 : Nettoyage des doublons ===
    unique_files = []
    seen = set()
    for f in files_data:
        if f["filename"] not in seen:
            unique_files.append(f)
            seen.add(f["filename"])

    conn.close()

    print(f"📎 {len(unique_files)} fichiers trouvés pour {cid} (dont {len(nouveaux_detectes)} nouveaux)")

    # (Sécurité) on ignore les statuts fantômes
    existing = {os.path.basename(f["path"]) for f in unique_files if os.path.exists(f["path"])}
    verif = {k: v for k, v in verif.items() if k in existing}

    return jsonify(unique_files)






@app.route("/admin/files/mark", methods=["POST"])
def admin_files_mark():
    if not require_admin():
        abort(403)

    data = request.json or {}
    cid = data.get("id")
    fname = data.get("filename")
    decision = data.get("decision")  # "conforme" ou "non_conforme"

    if not cid or not fname or decision not in ("conforme", "non_conforme"):
        return jsonify({"ok": False, "error": "paramètres invalides"}), 400

    conn = db()
    row = get_candidat(conn, cid)
    if not row:
        conn.close()
        return jsonify({"ok": False, "error": "candidat introuvable"}), 404

    verif = load_verif_docs(row)
    horodatage = datetime.now().strftime("%d/%m/%Y à %H:%M")

    # ✅ Marquer le document avec le label humain
    label_associe = ""
    for key, (field, label) in DOC_FIELDS.items():
        file_list = parse_list(row.get(field))
        if any(p.endswith(fname) for p in file_list):
            label_associe = label
            break

    verif[fname] = {
        "etat": decision,
        "horodatage": horodatage,
        "label": label_associe or "Pièce justificative"
    }

    cur = conn.cursor()

    # ❌ Si non conforme → supprimer physiquement le fichier
    if decision == "non_conforme":
        for key, (field, _) in DOC_FIELDS.items():
            file_list = parse_list(row.get(field))
            new_list = [p for p in file_list if not p.endswith(fname)]
            if len(new_list) != len(file_list):
                cur.execute(
                    f"UPDATE candidats SET {field}=?, updated_at=? WHERE id=?",
                    (json.dumps(new_list), datetime.now().isoformat(), cid)
                )
                try:
                    os.remove(os.path.join(UPLOAD_DIR, cid, fname))
                except FileNotFoundError:
                    pass

        # ✅ Mettre à jour le statut global
        cur.execute(
            "UPDATE candidats SET statut=?, verif_docs=?, updated_at=? WHERE id=?",
            ("docs_non_conformes", json.dumps(verif, ensure_ascii=False),
             datetime.now().isoformat(), cid)
        )

    else:
        # ✅ conforme → juste mise à jour
        cur.execute(
            "UPDATE candidats SET verif_docs=?, updated_at=? WHERE id=?",
            (json.dumps(verif, ensure_ascii=False), datetime.now().isoformat(), cid)
        )

    conn.commit()
    conn.close()

    log_event(row, "DOC_MARK", {"file": fname, "decision": decision})
    print(f"✅ Pièce {fname} marquée comme {decision}")
    return jsonify({"ok": True, "horodatage": horodatage})


# 💾 ROUTE : Fusionner les nouveaux fichiers dans les pièces normales
@app.route("/admin/files/merge", methods=["POST"])
def admin_files_merge():
    if not require_admin():
        abort(403)

    data = request.json or {}
    cid = data.get("id")
    if not cid:
        return jsonify({"ok": False, "error": "id manquant"}), 400

    conn = db()
    row = get_candidat(conn, cid)
    if not row:
        conn.close()
        return jsonify({"ok": False, "error": "candidat introuvable"}), 404

    # Charger le JSON meta existant
    try:
        meta = json.loads(row.get("replace_meta") or "{}")
        nouveaux = meta.get("nouveaux_fichiers", [])
        verif = load_verif_docs(row)
    except Exception:
        nouveaux = []
        verif = {}

    if not nouveaux:
        conn.close()
        return jsonify({"ok": False, "error": "aucun nouveau fichier"}), 400

    # Fusionner chaque nouveau fichier dans la bonne colonne
    for info in nouveaux:
        fname = info.get("fichier")
        label = info.get("label")
        if not fname:
            continue

        # Trouver la colonne correspondante selon le label
        for key, (col, label_ref) in DOC_FIELDS.items():
            if label_ref in label or label in label_ref:
                cur = conn.cursor()
                lst = parse_list(row.get(col))

                # 🧹 Supprime les anciens fichiers du même type
                prefix = fname.split("_")[0]
                lst = [p for p in lst if not os.path.basename(p).startswith(prefix)]

                # ➕ Ajoute la nouvelle version propre (avec le bon chemin)
                lst.append(os.path.join(UPLOAD_DIR, row["id"], fname))

                cur.execute(
                    f"UPDATE candidats SET {col}=?, updated_at=? WHERE id=?",
                    (json.dumps(lst, ensure_ascii=False), datetime.now().isoformat(), cid)
                )

                # 🧹 Remet le statut de ce fichier en "en_attente"
                verif[fname] = {
                    "etat": "en_attente",
                    "horodatage": datetime.now().strftime("%d/%m/%Y à %H:%M"),
                    "label": label
                }

                break

    # 🧩 Enregistrer le dictionnaire verif mis à jour et nettoyer replace_meta
    cur = conn.cursor()
    cur.execute(
        "UPDATE candidats SET nouveau_doc=0, replace_meta=?, verif_docs=?, updated_at=? WHERE id=?",
        ("{}", json.dumps(verif, ensure_ascii=False), datetime.now().isoformat(), cid)
    )
    conn.commit()

    # 🔄 Recharger la fiche pour nettoyage
    cur.execute("SELECT * FROM candidats WHERE id=?", (cid,))
    row = dict(cur.fetchone())

    # 🧽 Nettoyage des entrées orphelines dans verif_docs
    existing_files = []
    for key, (col, _) in DOC_FIELDS.items():
        existing_files += [os.path.basename(p) for p in parse_list(row.get(col))]

    verif_purged = {f: v for f, v in load_verif_docs(row).items() if f in existing_files}

    cur.execute(
        "UPDATE candidats SET verif_docs=?, updated_at=? WHERE id=?",
        (json.dumps(verif_purged, ensure_ascii=False), datetime.now().isoformat(), cid)
    )
    conn.commit()
    conn.close()

    return jsonify({"ok": True})


# =====================================================
# ✉️ NOTIFICATION DE DOCUMENTS NON CONFORMES
# =====================================================

@app.route("/admin/files/notify", methods=["POST"])
def admin_files_notify():
    if not require_admin():
        abort(403)

    data = request.json or {}
    cid = data.get("id")
    commentaire = (data.get("commentaire") or "").strip()

    if not cid:
        return jsonify({"ok": False, "error": "ID manquant"}), 400

    conn = db()
    row = get_candidat(conn, cid)
    if not row:
        conn.close()
        return jsonify({"ok": False, "error": "Candidat introuvable"}), 404

    verif = load_verif_docs(row)

    # 🔍 Lister les documents non conformes
    non_conformes = []
    for f, v in verif.items():
        if v.get("etat") == "non_conforme":
            label = v.get("label", "Pièce justificative")
            date = v.get("horodatage", "")
            non_conformes.append(f"{label} – {f} (le {date})")


    if not non_conformes:
        return jsonify({"ok": False, "error": "Aucune pièce non conforme"}), 400

    recap = "\n".join(["• " + n for n in non_conformes])

    # 🧩 Générer un token de remplacement valable 15 jours
    token = new_token()
    exp = (datetime.now() + timedelta(days=15)).isoformat()
    cur = conn.cursor()
    cur.execute("""
        UPDATE candidats 
        SET statut=?, replace_token=?, replace_token_exp=?, replace_meta=?, updated_at=? 
        WHERE id=?
    """, (
        "docs_non_conformes",
        token,
        exp,
        json.dumps({"pieces": non_conformes, "commentaire": commentaire}),
        datetime.now().isoformat(),
        cid
    ))
    conn.commit()
    conn.close()

    # ✉️ Envoi du mail au candidat
    link = make_signed_link("/replace-files", token)
    html = mail_html(
        "docs_non_conformes",
        prenom=row.get("prenom", ""),
        bts_label=BTS_LABELS.get((row.get("bts") or "").strip().upper(), row.get("bts")),
        lien_espace=link
    )
    send_mail(row.get("email", ""), "Documents non conformes – Intégrale Academy", html)

    # 📱 SMS documents non conformes
    tel = (row.get("tel", "") or "").replace(" ", "")
    if tel.startswith("0"):
        tel = "+33" + tel[1:]
    msg = sms_text(
        "docs_non_conformes",
        prenom=row.get("prenom", ""),
        bts_label=BTS_LABELS.get((row.get("bts") or "").strip().upper(), row.get("bts"))
    )
    send_sms_brevo(tel, msg)
    log_event(row, "SMS_ENVOYE", {"type": "docs_non_conformes", "tel": tel})

    log_event(row, "MAIL_ENVOYE", {"type": "docs_non_conformes", "pieces": non_conformes})

    return jsonify({"ok": True})

# =====================================================
# 📤 PAGE PUBLIQUE – RENVOI DE NOUVELLES PIÈCES
# =====================================================

@app.route("/replace-files", methods=["GET"])
def replace_files_form():
    token = request.args.get("token", "")
    sig = request.args.get("sig", "")
    if not verify_token(token, sig):
        abort(403)

    conn = db()
    cur = conn.cursor()
    cur.execute("SELECT * FROM candidats WHERE replace_token=?", (token,))
    row = cur.fetchone()
    conn.close()
    if not row:
        abort(404)

    row = dict(row)
    meta = json.loads(row.get("replace_meta") or "{}")
    pieces = meta.get("pieces", [])
    commentaire = meta.get("commentaire", "")

    return render_template(
        "replace_files.html",
        title="Envoyer mes nouvelles pièces justificatives",
        pieces=pieces,
        commentaire=commentaire,
        token=token,
        sig=sig
    )


@app.route("/replace-files", methods=["POST"])
def replace_files_submit():
    token = request.form.get("token", "")
    sig = request.form.get("sig", "")
    if not verify_token(token, sig):
        abort(403)

    conn = db()
    cur = conn.cursor()
    cur.execute("SELECT * FROM candidats WHERE replace_token=?", (token,))
    row = cur.fetchone()
    if not row:
        conn.close()
        abort(404)

    row = dict(row)

    # 🔁 Mise à jour du candidat : on sauvegarde les nouveaux fichiers dans replace_meta
    meta = json.loads(row.get("replace_meta") or "{}")
    nouveaux = []

    for field_name, (col_name, label) in DOC_FIELDS.items():
        files = request.files.getlist(field_name)
        prefix = FILE_PREFIX.get(field_name, field_name)

        # 📁 Dossier du candidat (même principe que save_files)
        cand_dir = os.path.join(UPLOAD_DIR, row["id"])
        os.makedirs(cand_dir, exist_ok=True)

        idx = 1
        for f in files:
            if not f or not f.filename:
                continue
            _, ext = os.path.splitext(secure_filename(f.filename))
            base = f"{prefix}_{row['id']}{'' if idx == 1 else f'_{idx}'}{ext.lower()}"
            dest = os.path.join(cand_dir, base)
            f.save(dest)
            nouveaux.append({"fichier": base, "label": label})
            idx += 1

    meta["nouveaux_fichiers"] = nouveaux

    cur.execute("""
        UPDATE candidats
        SET nouveau_doc=1,
            replace_meta=?,
            updated_at=?,
            statut=?
        WHERE id=?
    """, (
        json.dumps(meta, ensure_ascii=False),
        datetime.now().isoformat(),
        "preinscription",
        row["id"]
    ))
    conn.commit()
        # 🧹 Suppression automatique du badge relance si le candidat a renvoyé ses documents
    try:
        cur.execute("UPDATE candidats SET last_relance=NULL WHERE id=?", (row["id"],))
        conn.commit()
        print(f"🧹 Badge relance supprimé pour {row['prenom']} {row['nom']} (nouveaux docs reçus)")
    except Exception as e:
        print("⚠️ Erreur suppression badge relance (replace-files):", e)

    conn.close()

    # ✉️ Mail à l’admin
    admin_html = render_template(
        "mail_new_docs_admin.html",
        numero=row.get("numero_dossier", ""),
        nom=row.get("nom", ""),
        prenom=row.get("prenom", ""),
        fichiers=[n["fichier"] for n in nouveaux]
    )
    from_addr = os.getenv("MAIL_FROM", "ecole@integraleacademy.com")
    send_mail(from_addr, f"[ADMIN] Nouvelles pièces déposées ({row.get('numero_dossier')})", admin_html)

    log_event(row, "DOCS_RENVOYES", {"files": [n["fichier"] for n in nouveaux]})

    return render_template("replace_ok.html", title="Merci", fichiers=[n["fichier"] for n in nouveaux])








# =====================================================
# 📦 ROUTE : Télécharger toutes les pièces justificatives (ZIP)
# =====================================================

@app.route("/admin/files/download/<cid>")
def admin_files_download(cid):
    if not require_admin():
        abort(403)
    import zipfile, io

    conn = db()
    row = get_candidat(conn, cid)
    if not row:
        abort(404)
    conn.close()

    buffer = io.BytesIO()
    with zipfile.ZipFile(buffer, "w", zipfile.ZIP_DEFLATED) as zipf:
        for key, (field, label) in DOC_FIELDS.items():
            file_list = parse_list(row.get(field))
            for path in file_list:
                if os.path.exists(path):
                    zipf.write(path, arcname=os.path.basename(path))
    buffer.seek(0)

    zip_name = f"pieces_{row.get('numero_dossier','')}.zip"
    return send_file(
        buffer,
        mimetype="application/zip",
        as_attachment=True,
        download_name=zip_name
    )




@app.route("/admin/status/<cid>")
def admin_status(cid):
    if not require_admin():
        abort(403)
    conn = db()
    row = get_candidat(conn, cid)
    conn.close()
    if not row:
        return jsonify({"ok": False, "error": "not found"}), 404
    return jsonify({"ok": True, "statut": row.get("statut")})


# ---------------- Confirmation ----------------

def verify_token(query_token, query_sig):
    if not query_token or not query_sig:
        return False
    return sign_token(query_token) == query_sig


@app.route("/confirm-inscription", methods=["GET", "POST"])
def confirm_inscription():

    # -------------------- GET --------------------
    if request.method == "GET":
        token = request.args.get("token", "")
        sig = request.args.get("sig", "")

        if not verify_token(token, sig):
            abort(403)

        conn = db()
        cur = conn.cursor()
        cur.execute("SELECT * FROM candidats WHERE token_confirm=?", (token,))
        row = cur.fetchone()

        if not row:
            abort(404)

        row = dict(row)

        # 🛑 L'utilisateur a déjà confirmé son inscription
        if row.get("statut") == "confirmee":
            return render_template(
                "confirm_inscription.html",
                title="Inscription déjà confirmée",
                row=row,
                token=token,
                sig=sig,
                deja=True
            )

        # 🎓 Nom complet du BTS
        bts_label = BTS_LABELS.get(
            (row.get("bts") or "").strip().upper(),
            row.get("bts")
        )

        return render_template(
            "confirm_inscription.html",
            title="Confirmer mon inscription",
            row=row,
            token=token,
            sig=sig,
            bts_label=bts_label
        )


    # -------------------- POST --------------------
    token = request.form.get("token", "")
    sig = request.form.get("sig", "")

    if not verify_token(token, sig):
        abort(403)

    c1 = request.form.get("c1") == "on"
    c2 = request.form.get("c2") == "on"
    c3 = request.form.get("c3") == "on"

    if not (c1 and c2 and c3):
        flash("Merci de cocher les 3 cases obligatoires.", "error")
        return redirect(request.referrer or url_for("index"))

    conn = db()
    cur = conn.cursor()
    cur.execute("SELECT * FROM candidats WHERE token_confirm=?", (token,))
    row = dict(cur.fetchone())

    # ➕ Mise à jour du statut
    cur.execute(
        "UPDATE candidats SET statut=?, updated_at=? WHERE id=?",
        ("confirmee", datetime.now().isoformat(), row["id"])
    )
    conn.commit()

    # 🧹 Suppression badge relance
    try:
        cur.execute("UPDATE candidats SET last_relance=NULL WHERE id=?", (row["id"],))
        conn.commit()
    except Exception as e:
        print("⚠️ Erreur suppression badge relance :", e)

    # ---------------------- MAILS ----------------------
    mode_raw = (row.get("mode") or "").lower()

    if "pres" in mode_raw:
        mode_label = "En présentiel (Puget sur Argens, Var)"
    else:
        mode_label = "100% en ligne à distance en visioconférence ZOOM"

    # 👉 Construire lien espace candidat
    BASE_URL = os.getenv("BASE_URL", "https://inscriptionsbts.onrender.com").rstrip("/")
    slug = row.get("slug_public")
    lien_espace = f"{BASE_URL}/espace/{slug}"

    html = mail_html(
        "inscription_confirmee",
        prenom=row.get("prenom", ""),
        bts_label=BTS_LABELS.get((row.get("bts") or "").strip().upper(), row.get("bts")),
        numero_dossier=row.get("numero_dossier", ""),
        form_nom=row.get("nom", ""),
        form_prenom=row.get("prenom", ""),
        form_email=row.get("email", ""),
        form_tel=row.get("tel", ""),
        form_mode_label=mode_label,
        lien_espace=lien_espace   # 👈 ESSENTIEL
    )
    send_mail(row.get("email", ""), "Inscription confirmée – Intégrale Academy", html)

    log_event(row, "MAIL_ENVOYE", {"type": "inscription_confirmee"})
    log_event(row, "STATUT_CHANGE", {"statut": "confirmee"})

    return render_template("confirm_ok.html", title="Inscription confirmée")





@app.route("/reconfirm")
def reconfirm():
    token = request.args.get("token",""); sig = request.args.get("sig","")
    if not verify_token(token, sig): abort(403)
    conn = db(); cur = conn.cursor()
    cur.execute("SELECT * FROM candidats WHERE token_reconfirm=?", (token,))
    row = cur.fetchone()
    if not row: abort(404)
    row = dict(row)
    cur.execute("UPDATE candidats SET statut=?, updated_at=? WHERE id=?", ("reconfirmee", datetime.now().isoformat(), row["id"]))
    conn.commit()

    # 📱 SMS reconfirmation validée
    tel = (row.get("tel", "") or "").replace(" ", "")
    if tel.startswith("0"):
        tel = "+33" + tel[1:]
    msg = sms_text(
    "reconfirmation_validee",
    prenom=row.get("prenom", ""),
    bts_label=BTS_LABELS.get((row.get("bts") or "").strip().upper(), row.get("bts"))
)

    send_sms_brevo(tel, msg)
    log_event(row, "SMS_ENVOYE", {"type": "reconfirmation_validee", "tel": tel})


    log_event(row, "STATUT_CHANGE", {"statut": "reconfirmee"})
    return render_template("reconfirm_ok.html", title="Merci ❤️")

# =====================================================
# 🔁 PAGE PUBLIQUE – Reconfirmation manuelle (avec bouton)
# =====================================================

@app.route("/reconfirm-page")
def reconfirm_page():
    token = request.args.get("token", "")
    sig = request.args.get("sig", "")
    if not verify_token(token, sig):
        abort(403)

    conn = db()
    cur = conn.cursor()
    cur.execute("SELECT * FROM candidats WHERE token_reconfirm=?", (token,))
    row = cur.fetchone()
    conn.close()
    if not row:
        abort(404)

    row = dict(row)
    prenom = row.get("prenom", "")
    bts_label = BTS_LABELS.get((row.get("bts") or "").strip().upper(), row.get("bts"))

    return render_template(
        "reconfirm.html",
        prenom=prenom,
        bts_label=bts_label,
        token=token,
        sig=sig,
        title="Reconfirmer mon inscription"
    )


@app.route("/reconfirm-validate", methods=["POST"])
def reconfirm_validate():
    token = request.form.get("token", "")
    sig = request.form.get("sig", "")
    if not verify_token(token, sig):
        abort(403)

    conn = db()
    cur = conn.cursor()
    cur.execute("SELECT * FROM candidats WHERE token_reconfirm=?", (token,))
    row = cur.fetchone()
    if not row:
        conn.close()
        abort(404)

    row = dict(row)
    cur.execute(
        "UPDATE candidats SET statut=?, date_reconfirmee=?, updated_at=? WHERE id=?",
        ("reconfirmee", datetime.now().isoformat(), datetime.now().isoformat(), row["id"])
    )
    conn.commit()

    # 🧹 Suppression automatique du badge relance après reconfirmation
    try:
        cur.execute("UPDATE candidats SET last_relance=NULL WHERE id=?", (row["id"],))
        conn.commit()
        print(f"🧹 Badge relance supprimé pour {row['prenom']} {row['nom']} (reconfirmation)")
    except Exception as e:
        print("⚠️ Erreur suppression badge relance (reconfirm-validate):", e)


    conn.close()

    # ✉️ Mail confirmation reconfirmation
    html = mail_html(
        "reconfirmation_validee",
        prenom=row.get("prenom", ""),
        bts_label=BTS_LABELS.get((row.get("bts") or "").strip().upper(), row.get("bts"))
    )
    send_mail(row.get("email", ""), "Reconfirmation validée ✅", html)

    # 📱 SMS reconfirmation validée
    tel = (row.get("tel", "") or "").replace(" ", "")
    if tel.startswith("0"):
        tel = "+33" + tel[1:]
    msg = sms_text(
        "reconfirmation_validee",
        prenom=row.get("prenom", ""),
        bts_label=BTS_LABELS.get((row.get("bts") or "").strip().upper(), row.get("bts"))
    )
    send_sms_brevo(tel, msg)

    log_event(row, "STATUT_CHANGE", {"statut": "reconfirmee"})
    log_event(row, "MAIL_ENVOYE", {"type": "reconfirmation_validee"})
    log_event(row, "SMS_ENVOYE", {"type": "reconfirmation_validee", "tel": tel})

    return render_template("reconfirm_ok.html", prenom=row.get("prenom", ""), bts_label=row.get("bts"))



@app.route("/logout")
def logout():
    session.clear()
    return redirect(url_for("login"))

# ------------------------------------------------------------
# 📊 API publique : /data.json
#    → Sert au tableau de bord principal pour afficher le nombre
#      de pré-inscriptions à traiter en temps réel.
# ------------------------------------------------------------
@app.route("/data.json")
def data_json():
    try:
        conn = db()
        cur = conn.cursor()
        cur.execute("SELECT statut FROM candidats")
        rows = [r[0] for r in cur.fetchall()]

        # 🔍 Comptage : uniquement les statuts "pré-inscription à traiter"
        a_traiter = len([s for s in rows if not s or s == "preinscription" or s.lower().startswith("pré")])

        payload = {"a_traiter": a_traiter, "total": len(rows)}
        headers = {
            "Content-Type": "application/json",
            "Access-Control-Allow-Origin": "*"
        }
        return json.dumps(payload, ensure_ascii=False), 200, headers

    except Exception as e:
        print("❌ Erreur /data.json :", e)
        return json.dumps({"error": str(e)}), 500, {"Access-Control-Allow-Origin": "*"}

# =====================================================
# 👁️ ROUTE : Prévisualiser une pièce justificative
# =====================================================
@app.route("/uploads/<path:filename>")
def preview_upload(filename):
    # 🧩 Normalisation du nom
    filename = os.path.basename(filename)

    # 🧩 Chemin direct (compatibilité)
    direct_path = os.path.join(UPLOAD_DIR, filename)
    if os.path.exists(direct_path):
        return send_file(direct_path)

    # 🔎 Recherche dans les sous-dossiers candidats
    for cid in os.listdir(UPLOAD_DIR):
        sub_path = os.path.join(UPLOAD_DIR, cid, filename)
        if os.path.exists(sub_path):
            return send_file(sub_path)

    print(f"⚠️ Fichier introuvable : {filename}")
    abort(404)


# =====================================================
# 🧹 NETTOYAGE AUTOMATIQUE DES DOSSIERS ORPHELINS
# =====================================================
def cleanup_orphan_folders():
    print("🧹 Vérification des dossiers orphelins dans /uploads...")
    conn = db()
    cur = conn.cursor()
    cur.execute("SELECT id FROM candidats")
    valid_ids = {r[0] for r in cur.fetchall()}
    conn.close()

    import shutil
    removed = 0

    for f in os.listdir(UPLOAD_DIR):
        full = os.path.join(UPLOAD_DIR, f)
        if os.path.isdir(full) and f not in valid_ids:
            try:
                shutil.rmtree(full)
                removed += 1
            except Exception as e:
                print(f"⚠️ Impossible de supprimer {f}: {e}")

    if removed:
        print(f"✅ {removed} dossier(s) orphelin(s) supprimé(s) du répertoire uploads.")
    else:
        print("✅ Aucun dossier orphelin trouvé.")


# Lancer le nettoyage une fois au démarrage
with app.app_context():
    cleanup_orphan_folders()

with app.app_context():
    ensure_souhaite_accompagnement_field()


print("🚀 Application Flask démarrée – gestion CNAPS & pièces justificatives OK")

@app.route("/admin/clear-db", methods=["POST"])
def admin_clear_db():
    if not require_admin():
        abort(403)

    conn = db()
    cur = conn.cursor()
    # Supprimer tous les candidats et logs
    cur.execute("DELETE FROM candidats")
    cur.execute("DELETE FROM logs")
    conn.commit()
    conn.close()

    # 🧹 Supprimer tous les sous-dossiers candidats dans /uploads
    try:
        import shutil
        for f in os.listdir(UPLOAD_DIR):
            full = os.path.join(UPLOAD_DIR, f)
            if os.path.isdir(full):
                shutil.rmtree(full)
            elif os.path.isfile(full):
                os.remove(full)
    except Exception as e:
        print("⚠️ Erreur suppression fichiers :", e)

    flash("Base de données et fichiers effacés ✅", "success")
    return redirect(url_for("admin"))

    # =====================================================
# 🧩 AJOUT DU SLUG PUBLIC (IDENTIFIANT UNIQUE)
# =====================================================
def ensure_slug_public():
    conn = db()
    cur = conn.cursor()
    cur.execute("PRAGMA table_info(candidats)")
    cols = [r[1] for r in cur.fetchall()]
    if "slug_public" not in cols:
        print("🧩 Ajout de la colonne slug_public...")
        cur.execute("ALTER TABLE candidats ADD COLUMN slug_public TEXT DEFAULT ''")
        conn.commit()

    # Génération automatique d’un slug pour les candidats sans identifiant public
    cur.execute("SELECT id, slug_public FROM candidats")
    rows = cur.fetchall()
    for cid, slug in rows:
        if not slug:
            new_slug = uuid.uuid4().hex[:10]
            cur.execute("UPDATE candidats SET slug_public=? WHERE id=?", (new_slug, cid))
    conn.commit()
    conn.close()

with app.app_context():
    ensure_slug_public()

    # =====================================================
# 👤 PAGE PUBLIQUE – ESPACE CANDIDAT
# =====================================================
@app.route("/espace/<slug>")
def espace_candidat(slug):
    conn = db()
    cur = conn.cursor()
    cur.execute("SELECT * FROM candidats WHERE slug_public=?", (slug,))
    row = cur.fetchone()
    conn.close()

    if not row:
        abort(404)

    row = dict(row)
    statut = (row.get("statut") or "").lower()
    commentaire = row.get("commentaires") or ""

    # === Nom complet de la formation ===
    bts_code = (row.get("bts") or "").strip().upper()
    bts_label = BTS_LABELS.get(bts_code, row.get("bts"))
    row["bts_label"] = bts_label

        # === Libellé du mode de formation ===
    mode = (row.get("mode") or "").lower()
    if "dist" in mode:
        row["mode_label"] = "💻 À distance 100% en ligne (visioconférence)"
    else:
        row["mode_label"] = "🏫 En présentiel à Puget-sur-Argens (Var, 83)"

    explications = {
        "preinscription": "Nous avons bien reçu votre Pré-inscription. Votre dossier est cours d’examen par notre équipe.",
        "validee": "Votre candidature est validée. Vous devez confirmer votre inscription depuis le lien que nous vous avons envoyé par mail et par SMS.",
        "confirmee": "Votre inscription est officiellement confirmée 🎓. Bienvenue à Intégrale Academy !",
        "reconf_en_cours": "Vous souhaitez toujours intégrer notre école ? Vous avez changé d'avis ? Veuillez re-confirmer votre inscription depuis le mail que vous avez reçu.",
        "reconfirmee": "Tout est OK :✅.",
        "docs_non_conformes": "Certains documents ont été jugés non conformes. Veuillez consulter votre e-mail pour les renvoyer.",
        "annulee": "Votre inscription a été annulée. Pour toute question, contactez notre équipe.",
    }

    explication_statut = explications.get(statut, "Le traitement de votre dossier est en cours.")

    return render_template(
        "espace_candidat.html",
        row=row,
        statut=statut,
        explication_statut=explication_statut,
        commentaire=commentaire
    )




    # =====================================================
# 👁️ LIEN DIRECT DEPUIS L’ADMIN VERS L’ESPACE CANDIDAT
# =====================================================
@app.route("/admin/candidat/<cid>/espace")
def admin_espace_candidat(cid):
    if not require_admin():
        abort(403)

    conn = db()
    cur = conn.cursor()
    cur.execute("SELECT slug_public FROM candidats WHERE id=?", (cid,))
    row = cur.fetchone()
    conn.close()

    if not row or not row["slug_public"]:
        abort(404)

    slug = row["slug_public"]
    return redirect(url_for("espace_candidat", slug=slug))

# =====================================================
# 🔁 RECONFIRMATION MANUELLE D'INSCRIPTION (depuis admin)
# =====================================================
@app.route("/admin/reconfirm/<cid>", methods=["POST"])
def admin_reconfirm(cid):
    if not require_admin():
        abort(403)

    conn = db()
    cur = conn.cursor()
    cur.execute("SELECT * FROM candidats WHERE id=?", (cid,))
    row = cur.fetchone()
    if not row:
        conn.close()
        return jsonify({"ok": False, "error": "Candidat introuvable"}), 404

    row = dict(row)

    # 🔹 Génération du token de reconfirmation
    token = new_token()
    exp = (datetime.now() + timedelta(days=15)).isoformat()
    cur.execute(
        "UPDATE candidats SET token_reconfirm=?, token_reconfirm_exp=?, statut=?, updated_at=? WHERE id=?",
        (token, exp, "reconf_en_cours", datetime.now().isoformat(), cid)
    )
    conn.commit()
    conn.close()

    # 🔗 Lien de reconfirmation (signé)
    link = make_signed_link("/reconfirm-page", token)

    # ✉️ Mail HTML (utilise ton template)
    html = mail_html(
        "reconfirmation",
        prenom=row.get("prenom", ""),
        bts_label=BTS_LABELS.get((row.get("bts") or "").strip().upper(), row.get("bts")),
        lien_espace=link
    )

    # 🔹 Envoi du mail
    send_mail(
        row.get("email", ""),
        "Reconfirmation de votre inscription – Intégrale Academy",
        html
    )

    # 📱 SMS de notification (optionnel)
    tel = (row.get("tel", "") or "").replace(" ", "")
    if tel.startswith("0"):
        tel = "+33" + tel[1:]
    msg = sms_text(
        "reconfirmation_demandee",
        prenom=row.get("prenom", ""),
        bts_label=BTS_LABELS.get((row.get("bts") or "").strip().upper(), row.get("bts"))
    )
    send_sms_brevo(tel, msg)

    # 🧾 Log
    log_event(row, "MAIL_ENVOYE", {"type": "reconfirmation"})
    log_event(row, "SMS_ENVOYE", {"type": "reconfirmation", "tel": tel})
    log_event(row, "STATUT_CHANGE", {"statut": "reconf_en_cours"})

    return jsonify({"ok": True})


# =====================================================
# 🌐 GESTION DU PORTAIL – OUVERT / FERMÉ
# =====================================================

PORTAL_FILE = os.path.join(DATA_DIR, "portal.json")

def get_portal_status():
    """Lit le statut du portail (open/closed + message)"""
    if not os.path.exists(PORTAL_FILE):
        return {"status": "open", "message": "Ouvert aux inscriptions"}
    try:
        with open(PORTAL_FILE, "r", encoding="utf-8") as f:
            return json.load(f)
    except Exception:
        return {"status": "open", "message": "Ouvert aux inscriptions"}

def set_portal_status(status, message):
    """Met à jour le statut du portail"""
    data = {"status": status, "message": message}
    with open(PORTAL_FILE, "w", encoding="utf-8") as f:
        json.dump(data, f, indent=2, ensure_ascii=False)

# =====================================================
# 🔐 GESTION PORTAIL (Ouvrir / Fermer les inscriptions)
# =====================================================

PORTAL_FILE = os.path.join(DATA_DIR, "portal.json")

def load_portal_status():
    """Charge l'état actuel du portail (ou valeurs par défaut)."""
    if os.path.exists(PORTAL_FILE):
        with open(PORTAL_FILE, "r", encoding="utf-8") as f:
            try:
                return json.load(f)
            except json.JSONDecodeError:
                pass
    return {"status": "open", "message": "", "comment": ""}


# =====================================================
# 🌐 API – GESTION DU PORTAIL (État ouvert/fermé)
# =====================================================

@app.route("/get_portal_status")
def api_get_portal_status():
    """
    🔎 Route API pour récupérer l’état actuel du portail
    (utilisé par l’admin.js pour afficher le statut et le message).
    """
    data = load_portal_status()
    return jsonify(data)


@app.route("/set_portal_status", methods=["POST"])
def api_set_portal_status():
    """
    💾 Route API pour modifier le statut du portail
    (open / closed) + message + commentaire.
    """
    data = request.get_json() or {}
    status = data.get("status", "open")
    message = data.get("message", "")
    comment = data.get("comment", "")  # 🗒️ Nouveau champ commentaire

    # Enregistre le tout dans /data/portal.json
    with open(PORTAL_FILE, "w", encoding="utf-8") as f:
        json.dump(
            {"status": status, "message": message, "comment": comment},
            f,
            ensure_ascii=False,
            indent=2
        )

    print(f"✅ Portail mis à jour → {status.upper()} ({message})")
    if comment:
        print(f"🗒️ Commentaire : {comment}")

    return jsonify({"ok": True, "status": status, "message": message, "comment": comment})


# =====================================================
# ✉️📱 ADMIN – RENVOI MAILS (mail + SMS) — VERSION CORRIGÉE
# =====================================================
@app.route("/admin/resend_mail_sms/<cid>", methods=["POST"])
def admin_resend_mail_sms(cid):
    try:
        data = request.get_json(force=True)
        action = data.get("action")

        conn = db()
        row = conn.execute("SELECT * FROM candidats WHERE id=?", (cid,)).fetchone()
        if not row:
            return jsonify(error="Candidat introuvable"), 404

        row = dict(row)

        BASE_URL = os.getenv("BASE_URL", "https://inscriptionsbts.onrender.com").rstrip("/")
        slug = row.get("slug_public") or ""
        prenom = row["prenom"]
        email = row["email"]
        tel = row["tel"]
        bts_label = row["bts"]

        # 🔗 lien principal vers l’espace candidat
        lien_espace = f"{BASE_URL}/espace/{slug}"

        # =====================================================
        # 🎯 1. RECONSTRUIRE LE CONTEXTE EXACTEMENT COMME AU PREMIER MAIL
        # =====================================================
        from mail_templates import get_mail_context, mail_html
        from sms_templates import sms_text

        ctx = get_mail_context(row, lien_espace=lien_espace)

        # Nom complet BTS
        ctx["bts_label"] = BTS_LABELS.get(ctx["bts_label"], ctx["bts_label"])

        # =====================================================
        # 🔐 2. AJOUTER le lien de confirmation si "candidature validée"
        # =====================================================
        if action == "candidature_validee":
            token = row.get("token_confirm")
            if token:
                lien_confirmation = f"{BASE_URL}/confirm-inscription?token={token}&sig={sign_token(token)}"
                ctx["lien_confirmation"] = lien_confirmation

        # =====================================================
        # 📨 3. CHOISIR LE BON TEMPLATE
        # =====================================================
        mapping_mail = {
            "candidature_validee": "candidature_validee",
            "inscription_confirmee": "inscription_confirmee",
            "reconfirmation": "reconfirmation",
            "reconfirmee": "reconfirmation_validee",
            "docs_non_conformes": "docs_non_conformes",
        }

        mapping_sms = {
            "candidature_validee": "candidature_validee",
            "inscription_confirmee": "inscription_confirmee",
            "reconfirmation": "reconfirmation_demandee",
            "reconfirmee": "reconfirmation_validee",
            "docs_non_conformes": "docs_non_conformes",
        }

        tpl_mail = mapping_mail.get(action)
        tpl_sms = mapping_sms.get(action)

        if not tpl_mail:
            return jsonify(error="Action non reconnue"), 400

        # =====================================================
        # ✉️ 4. MAIL IDENTIQUE À L’ORIGINAL
        # =====================================================
        html_content = mail_html(tpl_mail, **ctx)
        subject = f"Intégrale Academy – {tpl_mail.replace('_', ' ').capitalize()}"
        send_mail(email, subject, html_content)

        # =====================================================
        # 📱 5. SMS IDENTIQUE À L’ORIGINAL
        # =====================================================
        tel_formate = (tel or "").replace(" ", "")
        if tel_formate.startswith("0"):
            tel_formate = "+33" + tel_formate[1:]

        sms_msg = sms_text(tpl_sms, prenom=prenom, bts_label=bts_label, lien_espace=lien_espace)
        send_sms_brevo(tel_formate, sms_msg)

        # LOG
        log_event({"id": cid}, "RENVOI_MAIL_SMS", {"type": action, "email": email, "tel": tel_formate})

        return jsonify(ok=True)

    except Exception as e:
        print("❌ Erreur resend_mail_sms :", e)
        return jsonify(error=str(e)), 500



# =====================================================
# 🔔 RELANCES (mail + SMS) — VERSION CORRIGÉE
# =====================================================
@app.route("/admin/relance/<cid>", methods=["POST"])
def admin_relance(cid):
    try:
        data = request.get_json() or {}
        action = data.get("action", "").strip()

        con = db()
        cur = con.cursor()
        cur.execute("SELECT * FROM candidats WHERE id=?", (cid,))
        row = cur.fetchone()

        if not row:
            return jsonify({"ok": False, "error": "Candidat introuvable"})

        row = dict(row)

        prenom = row["prenom"]
        email = row["email"]
        tel = row["tel"]
        bts_label = row["bts"]

        BASE_URL = os.getenv("BASE_URL", "https://inscriptionsbts.onrender.com").rstrip("/")
        slug = row.get("slug_public") or ""

        # ================================================
        # 🔗 1. Déterminer le lien EXACT comme dans le mail initial
        # ================================================
        if action == "candidature_validee":
            token = row.get("token_confirm")
            lien_action = f"{BASE_URL}/confirm-inscription?token={token}&sig={sign_token(token)}"

        elif action == "reconfirmation":
            token = row.get("token_reconfirm")
            lien_action = f"{BASE_URL}/reconfirm-page?token={token}&sig={sign_token(token)}"

        elif action == "docs_non_conformes":
            token = row.get("replace_token")
            lien_action = f"{BASE_URL}/replace-files?token={token}&sig={sign_token(token)}"

        else:
            # Default → Espace candidat
            lien_action = f"{BASE_URL}/espace/{slug}"

        # ================================================
        # 🎯 2. Charger le même contexte que le mail original
        # ================================================
        from mail_templates import get_mail_context, mail_html
        from sms_templates import sms_text
        from utils import send_mail, send_sms_brevo

        ctx = get_mail_context(row, lien_espace=lien_action)

        # Nom complet BTS
        ctx["bts_label"] = BTS_LABELS.get(ctx["bts_label"], ctx["bts_label"])

        # Cas particulier : si relance candidature validée → lien de confirmation nécessaire
        if action == "candidature_validee":
            ctx["lien_confirmation"] = lien_action

        # ================================================
        # ✉️ 3. Mappage des templates
        # ================================================
        mapping = {
            "candidature_validee": "relance_candidature_validee",
            "reconfirmation": "relance_reconfirmation",
            "docs_non_conformes": "relance_docs_non_conformes",
        }

        mail_tpl = mapping.get(action)
        if not mail_tpl:
            return jsonify({"ok": False, "error": "Action relance inconnue"})

        # ================================================
        # ✉️ 4. MAIL IDENTIQUE AU MAIL ORIGINAL
        # ================================================
        mail_html_content = mail_html(mail_tpl, **ctx)

        mail_subject = {
            "candidature_validee": "Relance – Confirmez votre inscription BTS en alternance",
            "reconfirmation": "Relance – Reconfirmez votre inscription",
            "docs_non_conformes": "Relance – Documents à compléter",
        }.get(action, "Relance – Intégrale Academy")

        mail_id = send_mail(email, mail_subject, mail_html_content)

        # ================================================
        # 📱 5. SMS IDENTIQUE AU SMS ORIGINAL
        # ================================================
        tel_fmt = str(tel).strip().replace(" ", "")
        if tel_fmt.startswith("0"):
            tel_fmt = "+33" + tel_fmt[1:]

        sms_id = send_sms_brevo(
            tel_fmt,
            sms_text(mail_tpl, prenom=prenom, bts_label=bts_label, lien_espace=lien_action)
        )

        # ================================================
        # 📝 6. LOG + date dernière relance
        # ================================================
        cur.execute(
            "INSERT INTO logs (cid, type, payload, created_at) VALUES (?, ?, ?, datetime('now','localtime'))",
            (cid, "RELANCE_ENVOYEE", f"{action} / mail_id: {mail_id} / sms_id: {sms_id}")
        )
        con.commit()

        cur.execute(
            "UPDATE candidats SET last_relance=?, updated_at=? WHERE id=?",
            (datetime.now().isoformat(), datetime.now().isoformat(), cid)
        )
        con.commit()

        con.close()

        print(f"✅ Relance envoyée à {prenom} ({email})")

        return jsonify({"ok": True})

    except Exception as e:
        print(f"❌ Erreur RELANCE :", e)
        return jsonify({"ok": False, "error": str(e)})



# =====================================================
# 🧮 COMPTE LES CANDIDATS "Inscription confirmée"
# =====================================================
@app.route("/admin/count_confirmed")
def admin_count_confirmed():
    try:
        conn = db()
        count = conn.execute("SELECT COUNT(*) FROM candidats WHERE statut = 'confirmee'").fetchone()[0]
        conn.close()
        return jsonify(ok=True, count=count)
    except Exception as e:
        print("Erreur count_confirmed:", e)
        return jsonify(ok=False, error=str(e)), 500


# =====================================================
# 📢 ENVOI DE RECONFIRMATION À TOUS LES CANDIDATS "Inscription confirmée"
# =====================================================
@app.route("/admin/reconfirm_all", methods=["POST"])
def admin_reconfirm_all():
    try:
        conn = db()
        cur = conn.cursor()

        # 🔍 On sélectionne uniquement les candidats avec le statut "confirmee"
        rows = cur.execute("""
            SELECT id, prenom, email, tel, bts, slug_public
            FROM candidats 
            WHERE statut = 'confirmee'
        """).fetchall()

        if not rows:
            conn.close()
            return jsonify(ok=False, error="Aucun candidat avec le statut 'Inscription confirmée'."), 400

        BASE_URL = os.getenv("BASE_URL", "https://inscriptionsbts.onrender.com").rstrip("/")
        sent_count = 0

        for r in rows:
            cid = r["id"]
            prenom = r["prenom"]
            email = r["email"]
            tel = (r["tel"] or "").replace(" ", "")
            bts_label = r["bts"]
            slug = r["slug_public"]

            # 🔑 Génère un nouveau token de reconfirmation
            token = new_token()
            exp = (datetime.now() + timedelta(days=15)).isoformat()

            cur.execute("""
                UPDATE candidats
                SET token_reconfirm=?, token_reconfirm_exp=?, statut=?, updated_at=?
                WHERE id=?
            """, (token, exp, "reconf_en_cours", datetime.now().isoformat(), cid))
            conn.commit()

            # 🔗 Lien de reconfirmation (signé)
            lien_espace = make_signed_link("/reconfirm-page", token)

            # ✉️ Mail de reconfirmation
            subject = "Reconfirmation d’inscription – Intégrale Academy"
            html_content = mail_html("reconfirmation", prenom=prenom, bts_label=bts_label, lien_espace=lien_espace)
            send_mail(email, subject, html_content)

            # 📱 SMS de reconfirmation
            if tel.startswith("0"):
                tel = "+33" + tel[1:]
            sms_msg = sms_text("reconfirmation_demandee", prenom=prenom, bts_label=bts_label, lien_espace=lien_espace)
            send_sms_brevo(tel, sms_msg)

            log_event({"id": cid}, "MAIL_ENVOYE", {"type": "reconfirmation"})
            log_event({"id": cid}, "SMS_ENVOYE", {"type": "reconfirmation", "tel": tel})
            log_event({"id": cid}, "STATUT_CHANGE", {"statut": "reconf_en_cours"})

            sent_count += 1
            print(f"📤 Reconfirmation envoyée à {prenom} ({email})")

        conn.close()
        return jsonify(ok=True, sent=sent_count)

    except Exception as e:
        print("❌ Erreur send_reconfirmation_all:", e)
        return jsonify(error=str(e)), 500


# =====================================================
# 🔔 MONITORING AUTOMATIQUE DE L'API BREVO
# =====================================================
import threading
import time
import requests

LAST_BREVO_ALERT = 0  # timestamp de la dernière alerte envoyée


def send_alert_brevo(message):
    """Envoie une alerte par Gmail (indépendant de Brevo)"""
    global LAST_BREVO_ALERT

    # Anti-spam : pas plus d'une alerte toutes les 30 minutes
    if time.time() - LAST_BREVO_ALERT < 1800:
        print("⏳ Alerte déjà envoyée il y a moins de 30 minutes, on ignore.")
        return

    LAST_BREVO_ALERT = time.time()

    try:
        # 📌 On utilise Gmail SMTP ici, jamais Brevo
        from utils import send_mail_gmail

        html = f"""
        <h2>⚠️ ALERTE CRITIQUE – API BREVO</h2>
        <p>{message}</p>
        <p><strong>L'API Brevo ne répond plus.</strong></p>
        <p>Site : inscriptionsbts.onrender.com</p>
        <p>Heure : {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}</p>
        """

        send_mail_gmail(
            "gestion@integraleacademy.com",
            "⚠️ ALERTE CRITIQUE – Échec API Brevo",
            html
        )

        print("🚨 Alerte envoyée à gestion@integraleacademy.com via Gmail (OK)")

    except Exception as e:
        print("❌ Impossible d'envoyer l'alerte Gmail :", e)


def check_brevo_health():
    """Teste la disponibilité de l'API Brevo toutes les 5 minutes"""
    while True:
        try:
            print("🔎 Vérification API Brevo...")

            api_key = os.getenv("BREVO_API_KEY", "")

            headers = {
                "api-key": api_key,
                "Content-Type": "application/json"
            }

            # petit appel simple à l'API Brevo
            r = requests.get("https://api.brevo.com/v3/account", headers=headers, timeout=10)

            if r.status_code != 200:
                msg = f"API Brevo en erreur : {r.status_code} – {r.text}"
                print("❌", msg)
                send_alert_brevo(msg)
            else:
                print("🟢 API Brevo OK")

        except Exception as e:
            msg = f"Erreur critique Brevo : {str(e)}"
            print("❌", msg)
            send_alert_brevo(msg)

        time.sleep(300)  # 5 minutes


# 🚀 Démarrage du monitoring Brevo
monitor_thread = threading.Thread(target=check_brevo_health, daemon=True)
monitor_thread.start()

@app.route("/brevo/status")
def brevo_status():
    """Retourne l'état actuel de Brevo"""
    import requests
    api_key = os.getenv("BREVO_API_KEY", "")

    try:
        r = requests.get(
            "https://api.brevo.com/v3/account",
            headers={"api-key": api_key, "Content-Type": "application/json"},
            timeout=5
        )
        if r.status_code == 200:
            return {"status": "ok"}
        else:
            return {"status": "error", "code": r.status_code}
    except:
        return {"status": "error", "code": "exception"}

@app.route("/admin/stats")
def admin_stats():
    if not require_admin():
        abort(403)

    conn = db()
    cur = conn.cursor()

    def count(statut):
        cur.execute("SELECT COUNT(*) FROM candidats WHERE statut=?", (statut,))
        return cur.fetchone()[0]

    stats = {
        "preinscription": count("preinscription"),
        "validee": count("validee"),
        "confirmee": count("confirmee"),
        "reconf_en_cours": count("reconf_en_cours"),
        "reconfirmee": count("reconfirmee"),
        "annulee": count("annulee"),
        "docs_non_conformes": count("docs_non_conformes"),
    }

    conn.close()
    return jsonify(ok=True, stats=stats)



# =====================================================
# 📩 ENVOI MANUEL DU MAIL APS (bouton dans l’admin)
# =====================================================
@app.route("/admin/send_mail_aps/<cid>", methods=["POST"])
def admin_send_mail_aps(cid):
    if not require_admin():
        return jsonify({"ok": False, "error": "Non autorisé"}), 403

    conn = db()
    cur = conn.cursor()
    cur.execute("SELECT * FROM candidats WHERE id=?", (cid,))
    row = cur.fetchone()
    conn.close()

    if not row:
        return jsonify({"ok": False, "error": "Candidat introuvable"}), 404

    row = dict(row)

    try:
        html = mail_html(
            "demande_aps",
            prenom=row.get("prenom", ""),
            bts_label=BTS_LABELS.get((row.get("bts") or "").strip().upper(), row.get("bts")),
            aps_session=row.get("aps_session", ""),
            lien_espace="https://cnapsv5-1.onrender.com/"
        )

        send_mail(
            row.get("email", ""),
            "🛡️ Formation APS – Documents CNAPS à envoyer",
            html
        )

        log_event(row, "MAIL_ENVOYE", {"type": "aps"})

        return jsonify({"ok": True})

    except Exception as e:
        print("❌ Erreur envoi mail APS :", e)
        return jsonify({"ok": False, "error": str(e)}), 500

from flask import send_from_directory

@app.route('/sw.js')
def sw():
    return send_from_directory('.', 'sw.js', mimetype='application/javascript')

@app.route('/brevo-frame.html')
def brevo_frame():
    return send_from_directory('.', 'brevo-frame.html', mimetype='text/html')





if __name__ == "__main__":
    app.run(host="0.0.0.0", port=5000, debug=True)
